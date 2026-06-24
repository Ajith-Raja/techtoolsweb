"""
PDF Tools API

This module provides a FastAPI server with multiple PDF processing tools:
- Compress PDF
- Convert PDF to images
- Convert images to PDF
- Extract text from PDF
- Extract images from PDF
- Merge PDFs
- Split PDF
- Add password protection
- Remove password protection
- Add watermark
- Rotate pages
- Reorder pages
- Add page numbers
- Add headers/footers
- PDF to Word conversion

Includes real-time progress tracking via WebSockets.
"""

import os
import io
import json
import tempfile
import uuid
import time
import asyncio
import concurrent.futures
from typing import List, Dict, Optional, Any
from enum import Enum
import logging

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from fastapi import (
    FastAPI, File, UploadFile, Form, WebSocket, WebSocketDisconnect,
    HTTPException, BackgroundTasks, Query, Body, Depends
)
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import uvicorn

logging.basicConfig(
    filename="app.log", 
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

logger = logging.getLogger("uvicorn.error")

# Initialize FastAPI
app = FastAPI(title="PDF Tools API", description="API for various PDF manipulation tasks")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create temp directory for file operations
TEMP_DIR = tempfile.gettempdir()
os.makedirs(os.path.join(TEMP_DIR, "pdf_tools"), exist_ok=True)

# In-memory storage for active tasks
active_tasks: Dict[str, Dict[str, Any]] = {}

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, task_id: str, websocket: WebSocket):
        await websocket.accept()
        self.active_connections[task_id] = websocket

    def disconnect(self, task_id: str):
        if task_id in self.active_connections:
            del self.active_connections[task_id]

    async def send_progress(self, task_id: str, data: Dict[str, Any]):
        if task_id in self.active_connections:
            await self.active_connections[task_id].send_json(data)

manager = ConnectionManager()

# Helper function to get a temporary file path
def get_temp_file(prefix: str, suffix: str) -> str:
    return os.path.join(TEMP_DIR, "pdf_tools", f"{prefix}_{str(uuid.uuid4())}{suffix}")


def _hex_to_rgb(color_value: str) -> tuple[int, int, int]:
    normalized = (color_value or "").strip().lstrip("#")
    if len(normalized) != 6:
        return (255, 0, 0)

    try:
        return tuple(int(normalized[index:index + 2], 16) for index in (0, 2, 4))  # type: ignore[return-value]
    except ValueError:
        return (255, 0, 0)


def _parse_page_ranges(page_ranges: str, total_pages: int) -> List[int]:
    selected_pages: List[int] = []

    for token in [part.strip() for part in (page_ranges or "").split(",") if part.strip()]:
        if "-" in token:
            start_text, end_text = [part.strip() for part in token.split("-", 1)]
            if not start_text.isdigit() or not end_text.isdigit():
                raise ValueError(f"Invalid page range: {token}")

            start_page = int(start_text)
            end_page = int(end_text)
            if start_page < 1 or end_page < start_page:
                raise ValueError(f"Invalid page range: {token}")

            for page_number in range(start_page, end_page + 1):
                if page_number > total_pages:
                    raise ValueError(f"Page {page_number} is out of range")
                selected_pages.append(page_number - 1)
        else:
            if not token.isdigit():
                raise ValueError(f"Invalid page number: {token}")

            page_number = int(token)
            if page_number < 1 or page_number > total_pages:
                raise ValueError(f"Page {page_number} is out of range")
            selected_pages.append(page_number - 1)

    unique_pages = list(dict.fromkeys(selected_pages))
    if not unique_pages:
        raise ValueError("No pages were selected")

    return unique_pages


def _watermark_position_rect(page: fitz.Page, position: WatermarkPosition, width: float, height: float) -> fitz.Rect:
    margin_x = page.rect.width * 0.06
    margin_y = page.rect.height * 0.06

    if position == WatermarkPosition.CENTER:
        x0 = (page.rect.width - width) / 2
        y0 = (page.rect.height - height) / 2
    elif position == WatermarkPosition.TOP_LEFT:
        x0 = margin_x
        y0 = margin_y
    elif position == WatermarkPosition.TOP_RIGHT:
        x0 = page.rect.width - width - margin_x
        y0 = margin_y
    elif position == WatermarkPosition.BOTTOM_LEFT:
        x0 = margin_x
        y0 = page.rect.height - height - margin_y
    else:
        x0 = page.rect.width - width - margin_x
        y0 = page.rect.height - height - margin_y

    return fitz.Rect(x0, y0, x0 + width, y0 + height)


def _wrap_text_to_lines(text: str, font_size: float, max_width: float, font_name: str = "helv") -> List[str]:
    """Wrap text so it fits within the available PDF width."""
    wrapped_lines: List[str] = []

    for raw_line in text.splitlines() or [text]:
        words = raw_line.split()

        if not words:
            wrapped_lines.append("")
            continue

        current_line = ""

        for word in words:
            candidate = word if not current_line else f"{current_line} {word}"

            if fitz.get_text_length(candidate, fontname=font_name, fontsize=font_size) <= max_width:
                current_line = candidate
                continue

            if current_line:
                wrapped_lines.append(current_line)

            if fitz.get_text_length(word, fontname=font_name, fontsize=font_size) <= max_width:
                current_line = word
                continue

            chunk = ""
            for character in word:
                candidate_chunk = f"{chunk}{character}"
                if fitz.get_text_length(candidate_chunk, fontname=font_name, fontsize=font_size) <= max_width:
                    chunk = candidate_chunk
                else:
                    if chunk:
                        wrapped_lines.append(chunk)
                    chunk = character

            current_line = chunk

        if current_line:
            wrapped_lines.append(current_line)

    return wrapped_lines or [""]


def _docx_alignment_to_pdf(alignment_value: Optional[int]) -> int:
    """Map python-docx alignment values to PyMuPDF alignment constants."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return fitz.TEXT_ALIGN_LEFT

    if alignment_value == WD_ALIGN_PARAGRAPH.CENTER:
        return fitz.TEXT_ALIGN_CENTER
    if alignment_value == WD_ALIGN_PARAGRAPH.RIGHT:
        return fitz.TEXT_ALIGN_RIGHT
    return fitz.TEXT_ALIGN_LEFT


def _iter_docx_body_elements(document):
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield DocxTable(child, document)


def _extract_run_images(run, document):
    import re

    image_blobs = []
    xml = getattr(run._element, "xml", "") or ""
    for embed in re.findall(r'r:embed="([^"]+)"', xml):
        if not embed:
            continue

        related_part = document.part.related_parts.get(embed)
        if related_part is None:
            continue

        image_blobs.append(related_part.blob)

    return image_blobs


def _paragraph_prefix(paragraph) -> str:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if "List Bullet" in style_name or "Bullet" in style_name:
        return "• "
    if "List Number" in style_name or "Number" in style_name:
        return "1. "
    return ""


def _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin):
    if cursor_y + required_height <= page_height - margin:
        return page, cursor_y

    page = pdf_document.new_page(width=page_width, height=page_height)
    return page, margin


def _render_docx_paragraph(pdf_document, page, cursor_y, paragraph, document, page_width, page_height, margin, usable_width):
    paragraph_text = paragraph.text.strip()
    prefix = _paragraph_prefix(paragraph)
    if prefix and paragraph_text:
        paragraph_text = f"{prefix}{paragraph_text}"
    if not paragraph_text and not any(_extract_run_images(run, document) for run in paragraph.runs):
        return page, cursor_y + 10

    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    font_size = 11
    if style_name.startswith("Title"):
        font_size = 18
    elif style_name.startswith("Heading 1"):
        font_size = 16
    elif style_name.startswith("Heading 2"):
        font_size = 14
    elif style_name.startswith("Heading"):
        font_size = 13

    is_bold = any(getattr(run, "bold", False) for run in paragraph.runs if run.text.strip())
    font_name = "helvb" if is_bold or style_name.startswith("Heading") or style_name == "Title" else "helv"
    alignment = _docx_alignment_to_pdf(paragraph.alignment)

    left_indent = 0.0
    first_line_indent = 0.0
    try:
        if paragraph.paragraph_format.left_indent is not None:
            left_indent = max(0.0, float(paragraph.paragraph_format.left_indent.pt))
        if paragraph.paragraph_format.first_line_indent is not None:
            first_line_indent = float(paragraph.paragraph_format.first_line_indent.pt)
    except Exception:
        pass

    x_base = margin + min(max(left_indent, 0.0), usable_width * 0.5)
    effective_width = max(80.0, usable_width - max(left_indent, 0.0))

    lines = _wrap_text_to_lines(paragraph_text, font_size, effective_width, font_name)
    line_height = font_size * 1.35
    paragraph_spacing = font_size * 0.55

    image_blobs = []
    for run in paragraph.runs:
        image_blobs.extend(_extract_run_images(run, document))

    required_height = (len(lines) * line_height) + paragraph_spacing + (len(image_blobs) * (usable_width * 0.6 + 14))
    page, cursor_y = _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin)

    for line_index, line in enumerate(lines):
        if cursor_y + line_height > page_height - margin:
            page = pdf_document.new_page(width=page_width, height=page_height)
            cursor_y = margin

        text_width = fitz.get_text_length(line, fontname=font_name, fontsize=font_size)
        if alignment == fitz.TEXT_ALIGN_CENTER:
            x_position = margin + max(0, (usable_width - text_width) / 2)
        elif alignment == fitz.TEXT_ALIGN_RIGHT:
            x_position = margin + max(0, usable_width - text_width)
        else:
            x_position = x_base + (first_line_indent if line_index == 0 else 0)

        page.insert_text(
            (x_position, cursor_y + font_size),
            line,
            fontsize=font_size,
            fontname=font_name,
            color=(0, 0, 0),
        )
        cursor_y += line_height

    for blob in image_blobs:
        try:
            image = Image.open(io.BytesIO(blob))
            if image.mode not in ("RGB", "RGBA"):
                image = image.convert("RGB")

            max_image_width = usable_width * 0.75
            max_image_height = (page_height - (2 * margin)) * 0.45
            scale = min(max_image_width / image.width, max_image_height / image.height, 1.0)
            draw_width = image.width * scale
            draw_height = image.height * scale

            if cursor_y + draw_height > page_height - margin:
                page = pdf_document.new_page(width=page_width, height=page_height)
                cursor_y = margin

            image_x = margin + max(0, (usable_width - draw_width) / 2)
            image_y = cursor_y + 4
            rect = fitz.Rect(image_x, image_y, image_x + draw_width, image_y + draw_height)
            page.insert_image(rect, stream=blob)
            cursor_y = image_y + draw_height + 8
        except Exception:
            continue

    return page, cursor_y + paragraph_spacing


def _render_docx_table(pdf_document, page, cursor_y, table, page_width, page_height, margin, usable_width):
    from docx.shared import Length

    row_count = len(table.rows)
    col_count = len(table.columns) if table.columns else max((len(row.cells) for row in table.rows), default=0)
    if row_count == 0 or col_count == 0:
        return page, cursor_y + 8

    table_width = usable_width
    col_width = table_width / col_count
    cell_padding = 4
    estimated_row_height = 26

    required_height = row_count * estimated_row_height + 8
    page, cursor_y = _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin)

    for row in table.rows:
        row_top = cursor_y
        row_height = estimated_row_height

        for cell_index, cell in enumerate(row.cells):
            cell_x0 = margin + (cell_index * col_width)
            cell_rect = fitz.Rect(cell_x0, row_top, cell_x0 + col_width, row_top + row_height)

            cell_text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
            if not cell_text:
                cell_text = ""

            page.draw_rect(cell_rect, color=(0, 0, 0), width=0.5)
            if cell_text:
                page.insert_textbox(
                    fitz.Rect(cell_rect.x0 + cell_padding, cell_rect.y0 + cell_padding, cell_rect.x1 - cell_padding, cell_rect.y1 - cell_padding),
                    cell_text,
                    fontsize=9,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

        cursor_y = row_top + row_height

    return page, cursor_y + 8


def _convert_docx_to_pdf(input_path: str, output_path: str) -> None:
    """Create a readable PDF from a DOCX document."""
    from docx import Document

    document = Document(input_path)
    pdf_document = fitz.open()

    page_width = 595
    page_height = 842
    margin = 48
    usable_width = page_width - (margin * 2)
    page = pdf_document.new_page(width=page_width, height=page_height)
    cursor_y = margin

    for block in _iter_docx_body_elements(document):
        if block.__class__.__name__ == "Table":
            page, cursor_y = _render_docx_table(pdf_document, page, cursor_y, block, page_width, page_height, margin, usable_width)
        else:
            page, cursor_y = _render_docx_paragraph(pdf_document, page, cursor_y, block, document, page_width, page_height, margin, usable_width)

    pdf_document.save(output_path, garbage=4, deflate=True)
    pdf_document.close()

# Helper function for updating task progress
async def update_progress(task_id: str, progress: int, status: str = "processing", error: str = None):
    if task_id in active_tasks:
        active_tasks[task_id]["progress"] = progress
        active_tasks[task_id]["status"] = status
        active_tasks[task_id]["error"] = error
        
        await manager.send_progress(
            task_id, 
            {
                "task_id": task_id,
                "progress": progress,
                "status": status,
                "error": error
            }
        )

# PDF password model
class PdfPasswordModel(BaseModel):
    owner_password: str
    user_password: Optional[str] = None


def _build_pdf_permissions(
    allow_printing: bool,
    allow_copying: bool,
    allow_editing: bool,
    allow_annotating: bool,
) -> int:
    permissions = 0

    if allow_printing:
        permissions |= fitz.PDF_PERM_PRINT | fitz.PDF_PERM_PRINT_HQ

    if allow_copying:
        permissions |= fitz.PDF_PERM_COPY | fitz.PDF_PERM_ACCESSIBILITY

    if allow_editing:
        permissions |= fitz.PDF_PERM_MODIFY | fitz.PDF_PERM_ASSEMBLE

    if allow_annotating:
        permissions |= fitz.PDF_PERM_ANNOTATE

    return int(permissions)

# PDF page range model
class PageRange(BaseModel):
    start: int
    end: int

# PDF rotation enum
class Rotation(int, Enum):
    DEGREES_0 = 0
    DEGREES_90 = 90
    DEGREES_180 = 180
    DEGREES_270 = 270

# Page size enum
class PageSize(str, Enum):
    A4 = "a4"
    LETTER = "letter"
    LEGAL = "legal"
    TABLOID = "tabloid"

# Watermark position enum
class WatermarkPosition(str, Enum):
    CENTER = "center"
    TOP_LEFT = "top-left"
    TOP_RIGHT = "top-right"
    BOTTOM_LEFT = "bottom-left"
    BOTTOM_RIGHT = "bottom-right"
    TILE = "tile"

# WebSocket endpoint for progress updates
@app.websocket("/ws/pdf/{task_id}")
async def websocket_endpoint(websocket: WebSocket, task_id: str):
    await manager.connect(task_id, websocket)
    try:
        while True:
            await websocket.receive_text()  # Keep connection alive
    except WebSocketDisconnect:
        manager.disconnect(task_id)

# Helper function to create a new task
def create_task() -> str:
    task_id = str(uuid.uuid4())
    active_tasks[task_id] = {
        "id": task_id,
        "progress": 0,
        "status": "created",
        "result_file": None,
        "error": None
    }
    return task_id

# Cleanup task after completion
async def cleanup_task(task_id: str, delay: int = 600):  # 10 minutes default
    await asyncio.sleep(delay)
    if task_id in active_tasks:
        result_file = active_tasks[task_id].get("result_file")
        if result_file and os.path.exists(result_file):
            try:
                os.unlink(result_file)
            except:
                pass
        del active_tasks[task_id]

# 1. Compress PDF
@app.post("/compress-pdf/", description="Compress a PDF file")
async def compress_pdf(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    quality: int = Form(80, description="Compression quality (0-100)")
):
    task_id = create_task()
    background_tasks.add_task(process_compress_pdf, task_id, pdf_file, quality)
    return {"task_id": task_id}

async def process_compress_pdf(task_id: str, pdf_file: UploadFile, quality: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("compressed", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        input_size = len(content)
        
        # Try Method A: Native Lossless Optimization (Best for text/tickets)
        # This keeps text vectorized, clear, and selectable.
        buffer = io.BytesIO()
        pdf_document.save(
            buffer, 
            garbage=4, 
            deflate=True, 
            deflate_images=True, 
            deflate_fonts=True
        )
        optimized_size = buffer.getbuffer().nbytes
        
        # If native compression successfully shrinks the file, use it
        if optimized_size < (input_size * 0.85):
            with open(output_path, "wb") as f:
                f.write(buffer.getvalue())
            pdf_document.close()
            
            active_tasks[task_id]["result_file"] = output_path
            await update_progress(task_id, 100, "completed")
            asyncio.create_task(cleanup_task(task_id))
            return

        # Method B Fallback: Image-level compression (retains vectorized text and layout, compresses only embedded images)
        if hasattr(pdf_document, "rewrite_images"):
            # Use PyMuPDF's built-in lossless & lossy image rewriting
            # This is extremely fast and preserves layout perfectly.
            pdf_document.rewrite_images(
                dpi_threshold=150,
                dpi_target=150,
                quality=quality,
                lossy=True,
                lossless=True
            )
            # Save the optimized PDF
            pdf_document.save(output_path, garbage=4, deflate=True)
            pdf_document.close()
        else:
            # Fallback: Manually iterate and replace images with compressed JPEGs
            processed_xrefs = set()
            total_pages = len(pdf_document)
            
            for i, page in enumerate(pdf_document):
                image_list = page.get_images(full=True)
                for img_info in image_list:
                    xref = img_info[0]
                    if xref in processed_xrefs:
                        continue
                    processed_xrefs.add(xref)
                    
                    try:
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        img = Image.open(io.BytesIO(image_bytes))
                        if img.mode != "RGB":
                            img = img.convert("RGB")
                        
                        # Downscale if excessively large
                        max_size = 1200
                        if max(img.size) > max_size:
                            img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                            
                        img_bytes = io.BytesIO()
                        img.save(img_bytes, format="JPEG", quality=quality)
                        page.replace_image(xref, stream=img_bytes.getvalue())
                    except Exception:
                        pass
                
                progress = 10 + int(85 * (i + 1) / total_pages)
                await update_progress(task_id, progress)
                
            pdf_document.save(output_path, garbage=4, deflate=True)
            pdf_document.close()
        
        # Guard: If compressed file is larger than original, try fallback to native optimized or original
        compressed_size = os.path.getsize(output_path)
        if compressed_size >= input_size:
            if optimized_size < input_size:
                with open(output_path, "wb") as f:
                    f.write(buffer.getvalue())
            else:
                import shutil
                shutil.copyfile(input_path, output_path)
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 2. Convert PDF to Images
@app.post("/pdf-to-images/", description="Convert PDF pages to images")
async def pdf_to_images(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    dpi: int = Form(200, description="Image DPI (150-600)"),
    format: str = Form("png", description="Image format (png, jpg, webp)")
):
    task_id = create_task()
    background_tasks.add_task(process_pdf_to_images, task_id, pdf_file, dpi, format)
    return {"task_id": task_id}

async def process_pdf_to_images(task_id: str, pdf_file: UploadFile, dpi: int, format: str):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("images", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("images", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Convert each page to image
        image_files = []
        for i, page in enumerate(pdf_document):
            pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
            
            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Save as requested format
            img_path = os.path.join(output_dir, f"page_{i+1}.{format}")
            img.save(img_path, format=format.upper())
            image_files.append(img_path)
            
            # Update progress
            progress = 10 + int(80 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Create ZIP archive with all images
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in image_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual image files
        for file in image_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 3. Convert Images to PDF
@app.post("/images-to-pdf/", description="Convert multiple images to a single PDF")
async def images_to_pdf(
    background_tasks: BackgroundTasks,
    images: List[UploadFile] = File(...),
    page_size: PageSize = Form(PageSize.A4),
    margin: int = Form(50, description="Margin in pixels")
):
    task_id = create_task()
    background_tasks.add_task(process_images_to_pdf, task_id, images, page_size, margin)
    return {"task_id": task_id}

async def process_images_to_pdf(task_id: str, images: List[UploadFile], page_size: PageSize, margin: int):
    output_path = get_temp_file("converted", ".pdf")
    temp_images = []
    
    try:
        # Page size dimensions (width, height) in points
        page_sizes = {
            PageSize.A4: (595, 842),
            PageSize.LETTER: (612, 792),
            PageSize.LEGAL: (612, 1008),
            PageSize.TABLOID: (792, 1224)
        }
        page_width, page_height = page_sizes[page_size]
        
        # Create new PDF
        pdf = fitz.open()
        
        total_images = len(images)
        for i, img_file in enumerate(images):
            # Save image to temp file
            img_path = get_temp_file(f"img_{i}", f".{img_file.filename.split('.')[-1]}")
            with open(img_path, "wb") as f:
                content = await img_file.read()
                f.write(content)
            temp_images.append(img_path)
            
            # Open image
            img = Image.open(img_path)
            
            # Calculate scaling to fit within page with margins
            img_width, img_height = img.size
            max_width = page_width - 2 * margin
            max_height = page_height - 2 * margin
            
            scale = min(max_width / img_width, max_height / img_height)
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            
            # Create new page
            page = pdf.new_page(width=page_width, height=page_height)
            
            # Calculate position to center image
            x = (page_width - new_width) / 2
            y = (page_height - new_height) / 2
            
            # Insert image
            page.insert_image(fitz.Rect(x, y, x + new_width, y + new_height), filename=img_path)
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_images)
            await update_progress(task_id, progress)
        
        # Save PDF
        pdf.save(output_path)
        pdf.close()
        
        # Cleanup temp images
        for img_path in temp_images:
            os.unlink(img_path)
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        for img_path in temp_images:
            if os.path.exists(img_path):
                os.unlink(img_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 4. Extract Text from PDF
@app.post("/extract-text/", description="Extract text from a PDF file")
async def extract_text(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_range: Optional[str] = Form(None, description="Page range (e.g., 1-5)")
):
    task_id = create_task()
    background_tasks.add_task(process_extract_text, task_id, pdf_file, page_range)
    return {"task_id": task_id}

async def process_extract_text(task_id: str, pdf_file: UploadFile, page_range: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("extracted_text", ".txt")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Parse page range
        if page_range:
            parts = page_range.split('-')
            start_page = max(0, int(parts[0]) - 1)  # Convert from 1-based to 0-based
            end_page = min(total_pages, int(parts[1]) if len(parts) > 1 else start_page + 1)
        else:
            start_page = 0
            end_page = total_pages
        
        # Extract text from pages
        with open(output_path, "w", encoding="utf-8") as f:
            for i in range(start_page, end_page):
                page = pdf_document[i]
                text = page.get_text()
                f.write(f"--- Page {i+1} ---\n\n")
                f.write(text)
                f.write("\n\n")
                
                # Update progress
                progress = 10 + int(85 * (i - start_page + 1) / (end_page - start_page))
                await update_progress(task_id, progress)
        
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 5. Extract Images from PDF
@app.post("/extract-images/", description="Extract images from a PDF file")
async def extract_images(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    min_size: int = Form(100, description="Minimum image size in pixels"),
    image_type: str = Form("png", description="Output image type: png or jpeg")
):
    task_id = create_task()
    normalized_image_type = (image_type or "png").lower()
    if normalized_image_type in ("jpg", "jpeg"):
        normalized_image_type = "jpeg"
    elif normalized_image_type != "png":
        normalized_image_type = "png"

    background_tasks.add_task(process_extract_images, task_id, pdf_file, min_size, normalized_image_type)
    return {"task_id": task_id}

async def process_extract_images(task_id: str, pdf_file: UploadFile, min_size: int, image_type: str):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("images", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("extracted_images", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Extract images from each page (xref images + inline content images)
        image_files = []
        image_count = 0
        output_extension = "png" if image_type == "png" else "jpg"
        seen_fingerprints = set()

        def _save_image_bytes(raw_bytes: bytes, min_dimension: int) -> bool:
            nonlocal image_count

            if not raw_bytes:
                return False

            fingerprint = (len(raw_bytes), hash(raw_bytes[:128]))
            if fingerprint in seen_fingerprints:
                return False

            try:
                pil_image = Image.open(io.BytesIO(raw_bytes))
                if min(pil_image.size) < min_dimension:
                    return False

                if image_type == "jpeg":
                    if pil_image.mode not in ("RGB", "L"):
                        pil_image = pil_image.convert("RGB")
                else:
                    if pil_image.mode not in ("RGB", "RGBA", "L"):
                        pil_image = pil_image.convert("RGBA")

                image_output = io.BytesIO()
                if image_type == "png":
                    pil_image.save(image_output, format="PNG")
                else:
                    pil_image.save(image_output, format="JPEG", quality=92, optimize=True)
                image_output.seek(0)

                img_path = os.path.join(output_dir, f"image_{image_count}.{output_extension}")
                with open(img_path, "wb") as f:
                    f.write(image_output.getvalue())

                image_files.append(img_path)
                image_count += 1
                seen_fingerprints.add(fingerprint)
                return True
            except Exception:
                return False

        for i, page in enumerate(pdf_document):
            image_list = page.get_images(full=True)

            # 1) Standard xref image extraction
            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = pdf_document.extract_image(xref)
                    _save_image_bytes(base_image.get("image", b""), min_size)
                except Exception:
                    continue

            # 2) Inline/content images that are not always returned by get_images
            try:
                blocks = page.get_text("dict").get("blocks", [])
                for block in blocks:
                    if block.get("type") == 1 and block.get("image"):
                        _save_image_bytes(block.get("image", b""), min_size)
            except Exception:
                pass

            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)

        # If strict min-size removed all images, retry once with no min filter.
        if image_count == 0 and min_size > 1:
            for i, page in enumerate(pdf_document):
                image_list = page.get_images(full=True)

                for img_info in image_list:
                    xref = img_info[0]
                    try:
                        base_image = pdf_document.extract_image(xref)
                        _save_image_bytes(base_image.get("image", b""), 1)
                    except Exception:
                        continue

                try:
                    blocks = page.get_text("dict").get("blocks", [])
                    for block in blocks:
                        if block.get("type") == 1 and block.get("image"):
                            _save_image_bytes(block.get("image", b""), 1)
                except Exception:
                    pass
        
        if image_count == 0:
            raise Exception("No images found in this PDF. Try lowering minimum size or use a PDF that contains embedded raster images.")

        # Create ZIP archive with all images
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in image_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual image files
        for file in image_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 6. Merge PDFs
@app.post("/merge-pdfs/", description="Merge multiple PDF files into one")
async def merge_pdfs(
    background_tasks: BackgroundTasks,
    pdf_files: List[UploadFile] = File(...)
):
    task_id = create_task()
    background_tasks.add_task(process_merge_pdfs, task_id, pdf_files)
    return {"task_id": task_id}

async def process_merge_pdfs(task_id: str, pdf_files: List[UploadFile]):
    temp_files = []
    output_path = get_temp_file("merged", ".pdf")
    
    try:
        await update_progress(task_id, 5, "processing")
        
        # Create merged PDF
        merged_pdf = fitz.open()
        
        total_files = len(pdf_files)
        for i, pdf_file in enumerate(pdf_files):
            # Save uploaded file
            input_path = get_temp_file(f"input_{i}", ".pdf")
            with open(input_path, "wb") as f:
                content = await pdf_file.read()
                f.write(content)
            temp_files.append(input_path)
            
            # Open PDF
            pdf_document = fitz.open(input_path)
            
            # Append to merged PDF
            merged_pdf.insert_pdf(pdf_document)
            pdf_document.close()
            
            # Update progress
            progress = 5 + int(90 * (i + 1) / total_files)
            await update_progress(task_id, progress)
        
        # Save merged PDF
        merged_pdf.save(output_path)
        merged_pdf.close()
        
        # Cleanup temp files
        for file in temp_files:
            os.unlink(file)
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        for file in temp_files:
            if os.path.exists(file):
                os.unlink(file)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 7. Split PDF
@app.post("/split-pdf/", description="Split a PDF into multiple files")
async def split_pdf(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    split_method: str = Form("pages", description="Split method: 'pages' or 'ranges'"),
    pages_per_pdf: Optional[int] = Form(None, description="Pages per output PDF"),
    page_ranges: Optional[str] = Form(None, description="Page ranges (e.g., '1-3,4-6,7-9')")
):
    task_id = create_task()
    background_tasks.add_task(process_split_pdf, task_id, pdf_file, split_method, pages_per_pdf, page_ranges)
    return {"task_id": task_id}

async def process_split_pdf(task_id: str, pdf_file: UploadFile, split_method: str, pages_per_pdf: Optional[int], page_ranges: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("split", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("split_pdfs", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Split based on method
        output_files = []
        
        if split_method == "pages" and pages_per_pdf:
            # Split by number of pages per PDF
            for i in range(0, total_pages, pages_per_pdf):
                # Calculate end page
                end = min(i + pages_per_pdf, total_pages)
                
                # Create new PDF
                output_file = os.path.join(output_dir, f"split_{i+1}-{end}.pdf")
                new_pdf = fitz.open()
                
                # Add pages from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=i, to_page=end-1)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + pages_per_pdf) / total_pages)
                await update_progress(task_id, min(95, progress))
        
        elif split_method == "ranges" and page_ranges:
            # Split by specified page ranges
            ranges = page_ranges.split(',')
            
            for i, range_str in enumerate(ranges):
                parts = range_str.strip().split('-')
                
                # Parse start and end pages (convert to 0-based)
                start = max(0, int(parts[0]) - 1)
                end = min(total_pages, int(parts[1]) if len(parts) > 1 else start + 1)
                
                # Create new PDF
                output_file = os.path.join(output_dir, f"split_{start+1}-{end}.pdf")
                new_pdf = fitz.open()
                
                # Add pages from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=start, to_page=end-1)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / len(ranges))
                await update_progress(task_id, min(95, progress))
        
        else:
            # Split into individual pages
            for i in range(total_pages):
                # Create new PDF with single page
                output_file = os.path.join(output_dir, f"page_{i+1}.pdf")
                new_pdf = fitz.open()
                
                # Add page from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=i, to_page=i)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / total_pages)
                await update_progress(task_id, min(95, progress))
        
        pdf_document.close()
        
        # Create ZIP archive with all split PDFs
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual PDF files
        for file in output_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 8. Add Password Protection
@app.post("/add-password/", description="Add password protection to a PDF")
async def add_password(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    owner_password: str = Form(...),
    user_password: Optional[str] = Form(None),
    allow_printing: bool = Form(True),
    allow_copying: bool = Form(True),
    allow_editing: bool = Form(False),
    allow_annotating: bool = Form(True)
):
    task_id = create_task()
    password = PdfPasswordModel(
        owner_password=owner_password,
        user_password=user_password
    )
    background_tasks.add_task(
        process_add_password,
        task_id,
        pdf_file,
        password,
        allow_printing,
        allow_copying,
        allow_editing,
        allow_annotating,
    )
    return {"task_id": task_id}

async def process_add_password(
    task_id: str,
    pdf_file: UploadFile,
    password: PdfPasswordModel,
    allow_printing: bool = True,
    allow_copying: bool = True,
    allow_editing: bool = False,
    allow_annotating: bool = True,
):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("protected", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 20, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        
        # Set permissions
        perm = _build_pdf_permissions(
            allow_printing=allow_printing,
            allow_copying=allow_copying,
            allow_editing=allow_editing,
            allow_annotating=allow_annotating,
        )
        
        pdf_document.save(
            output_path,
            encryption=fitz.PDF_ENCRYPT_AES_256,  # strongest encryption
            owner_pw=password.owner_password,
            user_pw=password.user_password or "",
            permissions=perm
        )
        
        pdf_document.close()
        
        await update_progress(task_id, 90, "processing")
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 9. Remove Password Protection
@app.post("/remove-password/", description="Remove password protection from a PDF")
async def remove_password(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    password: str = Form(...)
):
    task_id = create_task()
    background_tasks.add_task(process_remove_password, task_id, pdf_file, password)
    return {"task_id": task_id}

async def process_remove_password(task_id: str, pdf_file: UploadFile, password: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("unprotected", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 20, "processing")
        
        # Open PDF with password
        pdf_document = fitz.open(input_path)
        
        # Check if PDF is encrypted
        if pdf_document.is_encrypted:
            # Try to authenticate with password
            if not pdf_document.authenticate(password):
                raise Exception("Incorrect password")
            
            # Save without encryption
            pdf_document.save(output_path)
        else:
            # No password to remove, just copy
            pdf_document.save(output_path)
        
        pdf_document.close()
        
        await update_progress(task_id, 90, "processing")
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 10. Add Watermark
@app.post("/add-watermark/", description="Add text or image watermark to a PDF")
async def add_watermark(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    watermark_type: str = Form("text", description="text or image"),
    watermark_text: Optional[str] = Form(None),
    watermark_image: Optional[UploadFile] = File(None),
    position: WatermarkPosition = Form(WatermarkPosition.CENTER),
    opacity: float = Form(0.3, description="Opacity (0.1-1.0)"),
    rotation: int = Form(45, description="Rotation angle in degrees"),
    font_size: int = Form(36, description="Text watermark font size"),
    font_color: str = Form("#FF0000", description="Text watermark color as hex"),
    image_scale: float = Form(0.25, description="Image watermark scale relative to page width")
):
    task_id = create_task()
    background_tasks.add_task(
        process_add_watermark,
        task_id,
        pdf_file,
        watermark_type,
        watermark_text,
        watermark_image,
        position,
        opacity,
        rotation,
        font_size,
        font_color,
        image_scale,
    )
    return {"task_id": task_id}

async def process_add_watermark(
    task_id: str,
    pdf_file: UploadFile,
    watermark_type: str,
    watermark_text: Optional[str],
    watermark_image: Optional[UploadFile],
    position: WatermarkPosition,
    opacity: float,
    rotation: int,
    font_size: int,
    font_color: str,
    image_scale: float,
):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("watermarked", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)

        watermark_mode = (watermark_type or "text").strip().lower()
        image_bytes = None
        image_size = None
        if watermark_mode == "image":
            if watermark_image is None:
                raise ValueError("Watermark image is required for image watermark mode")
            image_bytes = await watermark_image.read()
            if not image_bytes:
                raise ValueError("Watermark image is empty")

            prepared_image = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
            if opacity < 1:
                alpha = prepared_image.getchannel("A")
                alpha = alpha.point(lambda value: int(value * max(0.0, min(opacity, 1.0))))
                prepared_image.putalpha(alpha)
            image_size = prepared_image.size

            if rotation:
                prepared_image = prepared_image.rotate(rotation, expand=True)
                image_size = prepared_image.size

            prepared_buffer = io.BytesIO()
            prepared_image.save(prepared_buffer, format="PNG")
            image_bytes = prepared_buffer.getvalue()
        else:
            if not watermark_text:
                raise ValueError("Watermark text is required for text watermark mode")

        text_color = _hex_to_rgb(font_color)
        font_size_value = max(8, int(font_size))
        
        # Process each page
        for i, page in enumerate(pdf_document):
            if watermark_mode == "image" and image_bytes is not None:
                watermark_image_pil = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
                base_width = max(24, page.rect.width * max(0.05, min(image_scale, 0.5)))
                scale = base_width / watermark_image_pil.width
                draw_width = max(24, watermark_image_pil.width * scale)
                draw_height = max(24, watermark_image_pil.height * scale)

                if position == WatermarkPosition.TILE:
                    step_x = max(draw_width * 1.4, 48)
                    step_y = max(draw_height * 1.4, 48)
                    y = 0.0
                    while y < page.rect.height:
                        x = 0.0
                        while x < page.rect.width:
                            rect = fitz.Rect(x, y, x + draw_width, y + draw_height)
                            page.insert_image(rect, stream=image_bytes, overlay=True)
                            x += step_x
                        y += step_y
                else:
                    rect = _watermark_position_rect(page, position, draw_width, draw_height)
                    page.insert_image(rect, stream=image_bytes, overlay=True)
            else:
                # Create a transparent text watermark image so rotation/opacity work consistently.
                font_candidates = ["arial.ttf", "Arial.ttf"]
                pil_font = None
                for font_candidate in font_candidates:
                    try:
                        pil_font = ImageFont.truetype(font_candidate, font_size_value)
                        break
                    except Exception:
                        continue

                if pil_font is None:
                    pil_font = ImageFont.load_default()

                text = watermark_text or ""
                bbox_canvas = Image.new("RGBA", (1, 1), (255, 255, 255, 0))
                bbox_draw = ImageDraw.Draw(bbox_canvas)
                text_bbox = bbox_draw.textbbox((0, 0), text, font=pil_font)
                text_width = max(1, text_bbox[2] - text_bbox[0])
                text_height = max(1, text_bbox[3] - text_bbox[1])
                text_canvas = Image.new("RGBA", (text_width + 40, text_height + 40), (255, 255, 255, 0))
                text_draw = ImageDraw.Draw(text_canvas)
                rgba_color = (*text_color, int(max(0.0, min(opacity, 1.0)) * 255))
                text_draw.text((20, 20), text, font=pil_font, fill=rgba_color)

                rotated_text = text_canvas.rotate(rotation, expand=True)
                text_buffer = io.BytesIO()
                rotated_text.save(text_buffer, format="PNG")
                text_bytes = text_buffer.getvalue()

                draw_width, draw_height = rotated_text.size
                if position == WatermarkPosition.TILE:
                    step_x = max(draw_width * 1.6, 80)
                    step_y = max(draw_height * 1.6, 80)
                    y = 0.0
                    while y < page.rect.height:
                        x = 0.0
                        while x < page.rect.width:
                            rect = fitz.Rect(x, y, x + draw_width, y + draw_height)
                            page.insert_image(rect, stream=text_bytes, overlay=True)
                            x += step_x
                        y += step_y
                else:
                    rect = _watermark_position_rect(page, position, draw_width, draw_height)
                    page.insert_image(rect, stream=text_bytes, overlay=True)
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 11. Edit Metadata
@app.post("/edit-metadata/", description="Edit PDF metadata")
async def edit_metadata(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    author: Optional[str] = Form(None),
    subject: Optional[str] = Form(None),
    keywords: Optional[str] = Form(None),
    creator: Optional[str] = Form(None),
    producer: Optional[str] = Form(None),
):
    task_id = create_task()
    background_tasks.add_task(
        process_edit_metadata,
        task_id,
        pdf_file,
        title,
        author,
        subject,
        keywords,
        creator,
        producer,
    )
    return {"task_id": task_id}


async def process_edit_metadata(
    task_id: str,
    pdf_file: UploadFile,
    title: Optional[str],
    author: Optional[str],
    subject: Optional[str],
    keywords: Optional[str],
    creator: Optional[str],
    producer: Optional[str],
):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("metadata-edited", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(await pdf_file.read())

        await update_progress(task_id, 20, "processing")

        pdf_document = fitz.open(input_path)
        existing_metadata = pdf_document.metadata or {}

        # Update only user-editable core fields to avoid format-specific metadata
        # validation errors from auxiliary keys.
        editable_fields = (
            "title",
            "author",
            "subject",
            "keywords",
            "creator",
            "producer",
        )
        metadata = {key: str(existing_metadata.get(key, "") or "") for key in editable_fields}
        metadata.update({
            "title": title if title is not None else metadata.get("title", ""),
            "author": author if author is not None else metadata.get("author", ""),
            "subject": subject if subject is not None else metadata.get("subject", ""),
            "keywords": keywords if keywords is not None else metadata.get("keywords", ""),
            "creator": creator if creator is not None else metadata.get("creator", ""),
            "producer": producer if producer is not None else metadata.get("producer", ""),
        })

        pdf_document.set_metadata(metadata)

        # Some PDFs contain XMP/XML metadata that can override document info shown by viewers.
        # Removing stale XML metadata ensures updated fields are reflected consistently.
        try:
            if hasattr(pdf_document, "get_xml_metadata") and pdf_document.get_xml_metadata():
                pdf_document.del_xml_metadata()
        except Exception:
            pass

        pdf_document.save(output_path, garbage=4, deflate=True)
        pdf_document.close()

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 12. Remove Pages
@app.post("/remove-pages/", description="Remove selected pages from a PDF")
async def remove_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_ranges: str = Form(..., description="Comma-separated pages/ranges to remove, e.g. '1,3,5-7'"),
):
    task_id = create_task()
    background_tasks.add_task(process_remove_pages, task_id, pdf_file, page_ranges)
    return {"task_id": task_id}


async def process_remove_pages(task_id: str, pdf_file: UploadFile, page_ranges: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("pages-removed", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(await pdf_file.read())

        await update_progress(task_id, 20, "processing")

        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        pages_to_remove = _parse_page_ranges(page_ranges, total_pages)
        pages_to_keep = [page_index for page_index in range(total_pages) if page_index not in pages_to_remove]

        if not pages_to_keep:
            raise ValueError("At least one page must remain in the PDF")

        pdf_document.select(pages_to_keep)
        await update_progress(task_id, 75, "processing")

        pdf_document.save(output_path, garbage=4, deflate=True)
        pdf_document.close()

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)
 
# 13. Rotate Pages
@app.post("/rotate-pages/", description="Rotate pages in a PDF")
async def rotate_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    angle: Rotation = Form(Rotation.DEGREES_90),
    page_range: Optional[str] = Form(None, description="Page range (e.g., 1-5)"),
    page_rotations: Optional[str] = Form(None, description="JSON map of page numbers to rotation angles")
):
    task_id = create_task()
    background_tasks.add_task(process_rotate_pages, task_id, pdf_file, angle, page_range, page_rotations)
    return {"task_id": task_id}

async def process_rotate_pages(
    task_id: str,
    pdf_file: UploadFile,
    angle: Rotation,
    page_range: Optional[str],
    page_rotations: Optional[str]
):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("rotated", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)

        logger.info(page_rotations)

        if page_rotations:
            parsed_rotations = json.loads(page_rotations)
            if not isinstance(parsed_rotations, dict):
                raise ValueError("page_rotations must be a JSON object")

            page_items = list(parsed_rotations.items())
            if not page_items:
                raise ValueError("page_rotations cannot be empty")

            for idx, (page_num_raw, rotation_raw) in enumerate(page_items):
                page_idx = int(page_num_raw) - 1
                if page_idx < 0 or page_idx >= total_pages:
                    raise ValueError(f"Page number out of range: {page_num_raw}")

                rotation_value = int(rotation_raw) % 360
                if rotation_value not in (0, 90, 180, 270):
                    raise ValueError(
                        f"Invalid rotation angle for page {page_num_raw}: {rotation_raw}. Use 0, 90, 180, or 270"
                    )

                pdf_document[page_idx].set_rotation((360 - rotation_value) % 360)

                progress = 10 + int(85 * (idx + 1) / len(page_items))
                await update_progress(task_id, progress)
        else:
            # Parse page range
            if page_range:
                parts = page_range.split('-')
                start_page = max(0, int(parts[0]) - 1)  # Convert from 1-based to 0-based
                end_page = min(total_pages, int(parts[1]) if len(parts) > 1 else start_page + 1)
            else:
                start_page = 0
                end_page = total_pages

            # Rotate pages
            for i in range(start_page, end_page):
                pdf_document[i].set_rotation((360 - int(angle)) % 360)

                # Update progress
                progress = 10 + int(85 * (i - start_page + 1) / (end_page - start_page))
                await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 12. Reorder Pages
@app.post("/reorder-pages/", description="Reorder pages in a PDF")
async def reorder_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_order: str = Form(..., description="Comma-separated page numbers (e.g., '3,1,2,4')")
):
    task_id = create_task()
    background_tasks.add_task(process_reorder_pages, task_id, pdf_file, page_order)
    return {"task_id": task_id}

async def process_reorder_pages(task_id: str, pdf_file: UploadFile, page_order: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("reordered", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Parse page order (convert from 1-based to 0-based)
        try:
            page_numbers = [int(p.strip()) - 1 for p in page_order.split(',')]
            
            # Validate page numbers
            if any(p < 0 or p >= total_pages for p in page_numbers):
                raise Exception(f"Invalid page numbers. Must be between 1 and {total_pages}.")
        except ValueError:
            raise Exception("Invalid page order format. Use comma-separated page numbers.")
        
        # Create new PDF with reordered pages
        new_pdf = fitz.open()
        
        for i, page_num in enumerate(page_numbers):
            new_pdf.insert_pdf(pdf_document, from_page=page_num, to_page=page_num)
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / len(page_numbers))
            await update_progress(task_id, progress)
        
        # Save result
        new_pdf.save(output_path)
        new_pdf.close()
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 13. Add Page Numbers
@app.post("/add-page-numbers/", description="Add page numbers to a PDF")
async def add_page_numbers(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    position: str = Form("bottom-right", description="Position: 'bottom-right', 'bottom-center', 'bottom-left', 'top-right', 'top-center', 'top-left'"),
    start_number: int = Form(1, description="Starting page number"),
    font_size: int = Form(12, description="Font size")
):
    task_id = create_task()
    background_tasks.add_task(process_add_page_numbers, task_id, pdf_file, position, start_number, font_size)
    return {"task_id": task_id}

async def process_add_page_numbers(task_id: str, pdf_file: UploadFile, position: str, start_number: int, font_size: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("page_numbers", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Add page numbers to each page
        for i in range(total_pages):
            page = pdf_document[i]
            page_number = start_number + i
            
            # Determine position
            margin = 30  # margin in points
            
            if position == "bottom-right":
                x = page.rect.width - margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_RIGHT
            elif position == "bottom-center":
                x = page.rect.width / 2
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_CENTER
            elif position == "bottom-left":
                x = margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_LEFT
            elif position == "top-right":
                x = page.rect.width - margin
                y = margin
                align = fitz.TEXT_ALIGN_RIGHT
            elif position == "top-center":
                x = page.rect.width / 2
                y = margin
                align = fitz.TEXT_ALIGN_CENTER
            elif position == "top-left":
                x = margin
                y = margin
                align = fitz.TEXT_ALIGN_LEFT
            else:
                # Default to bottom right
                x = page.rect.width - margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_RIGHT
            
            # Insert page number
            page.insert_text(
                (x, y),
                str(page_number),
                fontsize=font_size,
                align=align
            )
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 14. Add Headers/Footers
@app.post("/add-header-footer/", description="Add headers and/or footers to a PDF")
async def add_header_footer(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    header_text: Optional[str] = Form(None, description="Header text"),
    footer_text: Optional[str] = Form(None, description="Footer text"),
    font_size: int = Form(10, description="Font size")
):
    task_id = create_task()
    background_tasks.add_task(process_add_header_footer, task_id, pdf_file, header_text, footer_text, font_size)
    return {"task_id": task_id}

async def process_add_header_footer(task_id: str, pdf_file: UploadFile, header_text: Optional[str], footer_text: Optional[str], font_size: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("header_footer", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Add header/footer to each page
        for i in range(total_pages):
            page = pdf_document[i]
            
            # Add header if specified
            if header_text:
                page.insert_text(
                    (page.rect.width / 2, 20),  # Position at top center
                    header_text,
                    fontsize=font_size,
                    align=fitz.TEXT_ALIGN_CENTER
                )
            
            # Add footer if specified
            if footer_text:
                page.insert_text(
                    (page.rect.width / 2, page.rect.height - 20),  # Position at bottom center
                    footer_text,
                    fontsize=font_size,
                    align=fitz.TEXT_ALIGN_CENTER
                )
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 15. PDF to Word conversion
@app.post("/pdf-to-word/", description="Convert PDF to Word format (Note: Basic conversion only)")
async def pdf_to_word(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...)
):
    task_id = create_task()
    background_tasks.add_task(process_pdf_to_word, task_id, pdf_file)
    return {"task_id": task_id}

async def process_pdf_to_word(task_id: str, pdf_file: UploadFile):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("output", ".docx")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        try:
            # Try to use python-docx if available
            from docx import Document
            from docx.shared import Inches
            
            # Open PDF
            pdf_document = fitz.open(input_path)
            total_pages = len(pdf_document)
            
            # Create Word document
            doc = Document()
            
            # Process each page
            for i, page in enumerate(pdf_document):
                # Extract text
                text = page.get_text()
                
                # Add text to Word document
                doc.add_paragraph(text)
                
                # Add a page break except for the last page
                if i < total_pages - 1:
                    doc.add_page_break()
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / total_pages)
                await update_progress(task_id, progress)
            
            # Save Word document
            doc.save(output_path)
            
        except ImportError:
            # Fallback to simple text extraction if docx not available
            pdf_document = fitz.open(input_path)
            total_pages = len(pdf_document)
            
            # Extract text from all pages
            text = ""
            for i, page in enumerate(pdf_document):
                text += page.get_text()
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / total_pages)
                await update_progress(task_id, progress)
            
            # Save as text file with .docx extension
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)


# 16. Word to PDF conversion
@app.post("/word_to_pdf/", description="Convert Word documents to PDF format")
async def word_to_pdf(
    background_tasks: BackgroundTasks,
    word_file: UploadFile = File(...)
):
    task_id = create_task()
    file_bytes = await word_file.read()
    background_tasks.add_task(process_word_to_pdf, task_id, file_bytes)
    return {"task_id": task_id}


async def process_word_to_pdf(task_id: str, file_bytes: bytes):
    input_path = get_temp_file("input", ".docx")
    output_path = get_temp_file("output", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(file_bytes)

        await update_progress(task_id, 5, "processing")
        _convert_docx_to_pdf(input_path, output_path)

        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise Exception("Word to PDF conversion did not generate a PDF")

        await update_progress(task_id, 95, "processing")

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")

        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# Endpoint to get task status
@app.get("/task/{task_id}", description="Get the status of a task")
async def get_task_status(task_id: str):
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    return active_tasks[task_id]

# Endpoint to download the result file
@app.get("/download/{task_id}", description="Download the result file")
async def download_result(task_id: str):
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Task not completed")
    
    if not task["result_file"] or not os.path.exists(task["result_file"]):
        raise HTTPException(status_code=404, detail="Result file not found")
    
    # Get original filename from the path
    filename = os.path.basename(task["result_file"])

    lower_name = filename.lower()
    media_type = "application/octet-stream"
    if lower_name.endswith(".pdf"):
        media_type = "application/pdf"
    elif lower_name.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif lower_name.endswith(".txt"):
        media_type = "text/plain; charset=utf-8"
    elif lower_name.endswith(".zip"):
        media_type = "application/zip"
    
    return FileResponse(
        path=task["result_file"], 
        filename=filename,
        media_type=media_type
    )

# Health check endpoint
@app.get("/health", description="Health check endpoint")
async def health_check():
    return {"status": "healthy"}

# Run the server when executed directly
if __name__ == "__main__":
    uvicorn.run("pdf_tools_api:app", host="0.0.0.0", port=8001, reload=True)