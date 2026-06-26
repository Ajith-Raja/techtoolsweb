"""
PDF Tools API

This module provides a FastAPI server with multiple PDF processing tools:
- Compress PDF
- Convert PDF to images
- Convert images to PDF
- Extract text from PDF
- Extract images from PDF
- Merge PDFs
- Split PDF
- Add password protection
- Remove password protection
- Add watermark
- Rotate pages
- Reorder pages
- Add page numbers
- Add headers/footers
- PDF to Word conversion

Includes real-time progress tracking via WebSockets.
"""

import os
import io
import json
import tempfile
import uuid
import time
import asyncio
import concurrent.futures
from typing import List, Dict, Optional, Any
from enum import Enum
import subprocess
import shutil
import traceback
from datetime import datetime
import httpx
import redis

import fitz  # PyMuPDF
#import aspose.pdf as aspose_pdf
from PIL import Image, ImageDraw, ImageFont
from fastapi import (
    FastAPI, File, UploadFile, Form, WebSocket, WebSocketDisconnect,
    HTTPException, BackgroundTasks, Query, Body, Depends, Request
)
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from starlette.websockets import WebSocketState
import uvicorn
from yt_dlp import YoutubeDL
from sitemap_visualizer_api import analyze_sitemap
from image_to_svg_api import convert_image_to_svg
from qr_code_api import get_styles, generate_qr_code


# Initialize FastAPI
app = FastAPI(title="PDF Tools API", description="API for various PDF manipulation tasks")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://techtoolsweb.com", "http://localhost:5000"],  # In production, specify your frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create temp directory for file operations
TEMP_DIR = tempfile.gettempdir()
os.makedirs(os.path.join(TEMP_DIR, "pdf_tools"), exist_ok=True)

REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
TASK_TTL_SECONDS = int(os.getenv("TASK_TTL_SECONDS", "3600"))
TASK_STORE_PREFIX = os.getenv("TASK_STORE_PREFIX", "pdf_tools:task")


class TaskRecord(dict):
    def __init__(self, store: "RedisTaskStore", task_id: str, payload: Dict[str, Any]):
        super().__init__(payload)
        self._store = store
        self._task_id = task_id

    def _persist(self):
        self._store._save(self._task_id, dict(self))

    def __setitem__(self, key, value):
        super().__setitem__(key, value)
        self._persist()

    def __delitem__(self, key):
        super().__delitem__(key)
        self._persist()

    def update(self, *args, **kwargs):
        super().update(*args, **kwargs)
        self._persist()


class RedisTaskStore:
    def __init__(self, redis_url: str, prefix: str, ttl_seconds: int):
        self.prefix = prefix
        self.ttl_seconds = ttl_seconds
        self._redis_url = redis_url
        self._redis_client: Optional[redis.Redis] = None
        self._redis_unavailable = False
        self._fallback_store: Dict[str, Dict[str, Any]] = {}
        self._fallback_expiry: Dict[str, float] = {}

    def _key(self, task_id: str) -> str:
        return f"{self.prefix}:{task_id}"

    def _purge_fallback(self):
        now = time.time()
        expired = [task_id for task_id, expiry in self._fallback_expiry.items() if expiry <= now]
        for task_id in expired:
            self._fallback_store.pop(task_id, None)
            self._fallback_expiry.pop(task_id, None)

    def _client(self) -> Optional[redis.Redis]:
        if self._redis_unavailable:
            return None

        if self._redis_client is not None:
            return self._redis_client

        try:
            client = redis.Redis.from_url(self._redis_url, decode_responses=True)
            client.ping()
            self._redis_client = client
            return client
        except Exception as error:
            # Fall back to bounded local storage if Redis is unavailable.
            print(f"Redis unavailable, using in-process fallback store: {error}")
            self._redis_unavailable = True
            return None

    def _save(self, task_id: str, payload: Dict[str, Any]):
        client = self._client()
        if client is not None:
            client.set(self._key(task_id), json.dumps(payload), ex=self.ttl_seconds)
            return

        self._purge_fallback()
        self._fallback_store[task_id] = dict(payload)
        self._fallback_expiry[task_id] = time.time() + self.ttl_seconds

    def _load(self, task_id: str) -> Dict[str, Any]:
        client = self._client()
        if client is not None:
            value = client.get(self._key(task_id))
            if value is None:
                raise KeyError(task_id)
            client.expire(self._key(task_id), self.ttl_seconds)
            return json.loads(value)

        self._purge_fallback()
        if task_id not in self._fallback_store:
            raise KeyError(task_id)
        self._fallback_expiry[task_id] = time.time() + self.ttl_seconds
        return dict(self._fallback_store[task_id])

    def __contains__(self, task_id: str) -> bool:
        client = self._client()
        if client is not None:
            return bool(client.exists(self._key(task_id)))

        self._purge_fallback()
        return task_id in self._fallback_store

    def __getitem__(self, task_id: str) -> TaskRecord:
        payload = self._load(task_id)
        return TaskRecord(self, task_id, payload)

    def __setitem__(self, task_id: str, payload: Dict[str, Any]):
        self._save(task_id, payload)

    def __delitem__(self, task_id: str):
        client = self._client()
        if client is not None:
            client.delete(self._key(task_id))
            return

        self._fallback_store.pop(task_id, None)
        self._fallback_expiry.pop(task_id, None)

    def pop(self, task_id: str, default: Any = None):
        if task_id not in self:
            return default

        payload = self._load(task_id)
        self.__delitem__(task_id)
        return payload


# Redis-backed storage for active tasks with TTL to prevent memory leaks.
active_tasks = RedisTaskStore(REDIS_URL, TASK_STORE_PREFIX, TASK_TTL_SECONDS)

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, task_id: str, websocket: WebSocket):
        await websocket.accept()
        self.active_connections[task_id] = websocket

    def disconnect(self, task_id: str):
        self.active_connections.pop(task_id, None)

    #async def send_progress(self, task_id: str, data: Dict[str, Any]):
    #    if task_id in self.active_connections:
    #        await self.active_connections[task_id].send_json(data)

    async def send_progress(self, task_id: str, data: Dict[str, Any]):
        websocket = self.active_connections.get(task_id)
        if not websocket:
            return

        if (
            websocket.application_state != WebSocketState.CONNECTED
            or websocket.client_state == WebSocketState.DISCONNECTED
        ):
            self.disconnect(task_id)
            return

        try:
            await websocket.send_json(data)
        except Exception as e:
            print(f"Error sending progress to task {task_id}: {e}")
            self.disconnect(task_id)

    async def shutdown(self):
        for task_id, websocket in list(self.active_connections.items()):
            try:
                if websocket.application_state == WebSocketState.CONNECTED:
                    await websocket.close(code=1001, reason="Server shutdown")
            except Exception:
                pass
            self.disconnect(task_id)

class InMemoryDB:
    def __init__(self):
        self.requests_history = []
        self.collections = {}
        self.environments = {}

db = InMemoryDB()

class HttpMethod(str, Enum):
    GET = "GET"
    POST = "POST"
    PUT = "PUT"
    PATCH = "PATCH"
    DELETE = "DELETE"
    OPTIONS = "OPTIONS"
    HEAD = "HEAD"

class BodyType(str, Enum):
    NONE = "none"
    RAW = "raw"
    FORM_DATA = "form-data"
    URLENCODED = "x-www-form-urlencoded"
    BINARY = "binary"

class RawBodyFormat(str, Enum):
    JSON = "json"
    XML = "xml"
    HTML = "html"
    TEXT = "text"

class ApiRequest(BaseModel):
    method: HttpMethod
    url: str
    headers: Optional[Dict[str, str]] = {}
    query_params: Optional[Dict[str, str]] = {}
    body_type: BodyType = BodyType.NONE
    raw_body_format: Optional[RawBodyFormat] = None
    body: Optional[Any] = None
    form_data: Optional[Dict[str, str]] = {}
    auth: Optional[Dict[str, str]] = {}

class SavedRequest(BaseModel):
    id: str
    name: str
    request: ApiRequest

class Collection(BaseModel):
    id: str
    name: str
    description: Optional[str] = None
    requests: List[SavedRequest] = []

class Environment(BaseModel):
    id: str
    name: str
    variables: Dict[str, str] = {}

class ApiResponse(BaseModel):
    request_id: str
    request: ApiRequest
    status_code: int
    headers: Dict[str, str]
    body: Any
    cookies: Dict[str, str]
    time_taken: float  # in milliseconds
    size: int  # response size in bytes
    timestamp: datetime

manager = ConnectionManager()

# Helper function to get a temporary file path
def get_temp_file(prefix: str, suffix: str) -> str:
    return os.path.join(TEMP_DIR, "pdf_tools", f"{prefix}_{str(uuid.uuid4())}{suffix}")


def _wrap_text_to_lines(text: str, font_size: float, max_width: float, font_name: str = "helv") -> List[str]:
    """Wrap text so it fits within the available PDF width."""
    wrapped_lines: List[str] = []

    for raw_line in text.splitlines() or [text]:
        words = raw_line.split()

        if not words:
            wrapped_lines.append("")
            continue

        current_line = ""

        for word in words:
            candidate = word if not current_line else f"{current_line} {word}"

            if fitz.get_text_length(candidate, fontname=font_name, fontsize=font_size) <= max_width:
                current_line = candidate
                continue

            if current_line:
                wrapped_lines.append(current_line)

            if fitz.get_text_length(word, fontname=font_name, fontsize=font_size) <= max_width:
                current_line = word
                continue

            chunk = ""
            for character in word:
                candidate_chunk = f"{chunk}{character}"
                if fitz.get_text_length(candidate_chunk, fontname=font_name, fontsize=font_size) <= max_width:
                    chunk = candidate_chunk
                else:
                    if chunk:
                        wrapped_lines.append(chunk)
                    chunk = character

            current_line = chunk

        if current_line:
            wrapped_lines.append(current_line)

    return wrapped_lines or [""]


def _docx_alignment_to_pdf(alignment_value: Optional[int]) -> int:
    """Map python-docx alignment values to PyMuPDF alignment constants."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return fitz.TEXT_ALIGN_LEFT

    if alignment_value == WD_ALIGN_PARAGRAPH.CENTER:
        return fitz.TEXT_ALIGN_CENTER
    if alignment_value == WD_ALIGN_PARAGRAPH.RIGHT:
        return fitz.TEXT_ALIGN_RIGHT
    return fitz.TEXT_ALIGN_LEFT


def _iter_docx_body_elements(document):
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield DocxTable(child, document)


def _extract_run_images(run, document):
    import re

    image_blobs = []
    xml = getattr(run._element, "xml", "") or ""
    for embed in re.findall(r'r:embed="([^"]+)"', xml):
        if not embed:
            continue

        related_part = document.part.related_parts.get(embed)
        if related_part is None:
            continue

        image_blobs.append(related_part.blob)

    return image_blobs


def _paragraph_prefix(paragraph) -> str:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if "List Bullet" in style_name or "Bullet" in style_name:
        return "• "
    if "List Number" in style_name or "Number" in style_name:
        return "1. "
    return ""


def _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin):
    if cursor_y + required_height <= page_height - margin:
        return page, cursor_y

    page = pdf_document.new_page(width=page_width, height=page_height)
    return page, margin


def _render_docx_paragraph(pdf_document, page, cursor_y, paragraph, document, page_width, page_height, margin, usable_width):
    paragraph_text = paragraph.text.strip()
    prefix = _paragraph_prefix(paragraph)
    if prefix and paragraph_text:
        paragraph_text = f"{prefix}{paragraph_text}"
    if not paragraph_text and not any(_extract_run_images(run, document) for run in paragraph.runs):
        return page, cursor_y + 10

    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    font_size = 11
    if style_name.startswith("Title"):
        font_size = 18
    elif style_name.startswith("Heading 1"):
        font_size = 16
    elif style_name.startswith("Heading 2"):
        font_size = 14
    elif style_name.startswith("Heading"):
        font_size = 13

    is_bold = any(getattr(run, "bold", False) for run in paragraph.runs if run.text.strip())
    font_name = "helvb" if is_bold or style_name.startswith("Heading") or style_name == "Title" else "helv"
    alignment = _docx_alignment_to_pdf(paragraph.alignment)

    left_indent = 0.0
    first_line_indent = 0.0
    try:
        if paragraph.paragraph_format.left_indent is not None:
            left_indent = max(0.0, float(paragraph.paragraph_format.left_indent.pt))
        if paragraph.paragraph_format.first_line_indent is not None:
            first_line_indent = float(paragraph.paragraph_format.first_line_indent.pt)
    except Exception:
        pass

    x_base = margin + min(max(left_indent, 0.0), usable_width * 0.5)
    effective_width = max(80.0, usable_width - max(left_indent, 0.0))

    lines = _wrap_text_to_lines(paragraph_text, font_size, effective_width, font_name)
    line_height = font_size * 1.35
    paragraph_spacing = font_size * 0.55

    image_blobs = []
    for run in paragraph.runs:
        image_blobs.extend(_extract_run_images(run, document))

    required_height = (len(lines) * line_height) + paragraph_spacing + (len(image_blobs) * (usable_width * 0.6 + 14))
    page, cursor_y = _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin)

    for line_index, line in enumerate(lines):
        if cursor_y + line_height > page_height - margin:
            page = pdf_document.new_page(width=page_width, height=page_height)
            cursor_y = margin

        text_width = fitz.get_text_length(line, fontname=font_name, fontsize=font_size)
        if alignment == fitz.TEXT_ALIGN_CENTER:
            x_position = margin + max(0, (usable_width - text_width) / 2)
        elif alignment == fitz.TEXT_ALIGN_RIGHT:
            x_position = margin + max(0, usable_width - text_width)
        else:
            x_position = x_base + (first_line_indent if line_index == 0 else 0)

        page.insert_text(
            (x_position, cursor_y + font_size),
            line,
            fontsize=font_size,
            fontname=font_name,
            color=(0, 0, 0),
        )
        cursor_y += line_height

    for blob in image_blobs:
        try:
            image = Image.open(io.BytesIO(blob))
            if image.mode not in ("RGB", "RGBA"):
                image = image.convert("RGB")

            max_image_width = usable_width * 0.75
            max_image_height = (page_height - (2 * margin)) * 0.45
            scale = min(max_image_width / image.width, max_image_height / image.height, 1.0)
            draw_width = image.width * scale
            draw_height = image.height * scale

            if cursor_y + draw_height > page_height - margin:
                page = pdf_document.new_page(width=page_width, height=page_height)
                cursor_y = margin

            image_x = margin + max(0, (usable_width - draw_width) / 2)
            image_y = cursor_y + 4
            rect = fitz.Rect(image_x, image_y, image_x + draw_width, image_y + draw_height)
            page.insert_image(rect, stream=blob)
            cursor_y = image_y + draw_height + 8
        except Exception:
            continue

    return page, cursor_y + paragraph_spacing


def _render_docx_table(pdf_document, page, cursor_y, table, page_width, page_height, margin, usable_width):
    row_count = len(table.rows)
    col_count = len(table.columns) if table.columns else max((len(row.cells) for row in table.rows), default=0)
    if row_count == 0 or col_count == 0:
        return page, cursor_y + 8

    table_width = usable_width
    col_width = table_width / col_count
    cell_padding = 4
    estimated_row_height = 26

    required_height = row_count * estimated_row_height + 8
    page, cursor_y = _ensure_page_space(pdf_document, page, cursor_y, required_height, page_width, page_height, margin)

    for row in table.rows:
        row_top = cursor_y
        row_height = estimated_row_height

        for cell_index, cell in enumerate(row.cells):
            cell_x0 = margin + (cell_index * col_width)
            cell_rect = fitz.Rect(cell_x0, row_top, cell_x0 + col_width, row_top + row_height)

            cell_text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
            if not cell_text:
                cell_text = ""

            page.draw_rect(cell_rect, color=(0, 0, 0), width=0.5)
            if cell_text:
                page.insert_textbox(
                    fitz.Rect(cell_rect.x0 + cell_padding, cell_rect.y0 + cell_padding, cell_rect.x1 - cell_padding, cell_rect.y1 - cell_padding),
                    cell_text,
                    fontsize=9,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

        cursor_y = row_top + row_height

    return page, cursor_y + 8


def _convert_docx_to_pdf(input_path: str, output_path: str) -> None:
    """Create a readable PDF from a DOCX document."""
    from docx import Document

    document = Document(input_path)
    pdf_document = fitz.open()

    page_width = 595
    page_height = 842
    margin = 48
    usable_width = page_width - (margin * 2)
    page = pdf_document.new_page(width=page_width, height=page_height)
    cursor_y = margin

    for block in _iter_docx_body_elements(document):
        if block.__class__.__name__ == "Table":
            page, cursor_y = _render_docx_table(pdf_document, page, cursor_y, block, page_width, page_height, margin, usable_width)
        else:
            page, cursor_y = _render_docx_paragraph(pdf_document, page, cursor_y, block, document, page_width, page_height, margin, usable_width)

    pdf_document.save(output_path, garbage=4, deflate=True)
    pdf_document.close()

def build_temp_file(file: str) -> str:
    return os.path.join(TEMP_DIR, "pdf_tools", f"{file}")

# Helper function for updating task progress
async def update_progress(
    task_id: str,
    progress: int,
    status: str = "processing",
    error: str = None,
    file: str = None,
    stage: str = None,
    error_details: str = None,
):
    if task_id in active_tasks:
        active_tasks[task_id]["progress"] = progress
        active_tasks[task_id]["status"] = status
        active_tasks[task_id]["error"] = error
        if stage is not None:
            active_tasks[task_id]["stage"] = stage
        if error_details is not None:
            active_tasks[task_id]["error_details"] = error_details

        await manager.send_progress(
            task_id, 
            {
                "task_id": task_id,
                "progress": progress,
                "status": status,
                "error": error,
                "stage": active_tasks[task_id].get("stage"),
                "error_details": active_tasks[task_id].get("error_details")
            }
        )

# PDF password model
class PdfPasswordModel(BaseModel):
    owner_password: str
    user_password: Optional[str] = None

# PDF page range model
class PageRange(BaseModel):
    start: int
    end: int

# PDF rotation enum
class Rotation(int, Enum):
    DEGREES_0 = 0
    DEGREES_90 = 90
    DEGREES_180 = 180
    DEGREES_270 = 270

# Page size enum
class PageSize(str, Enum):
    A4 = "a4"
    LETTER = "letter"
    LEGAL = "legal"
    TABLOID = "tabloid"

# Watermark position enum
class WatermarkPosition(str, Enum):
    CENTER = "center"
    TOP_LEFT = "top-left"
    TOP_RIGHT = "top-right"
    BOTTOM_LEFT = "bottom-left"
    BOTTOM_RIGHT = "bottom-right"

# WebSocket endpoint for progress updates
@app.websocket("/ws/pdf/{task_id}")
async def websocket_endpoint(websocket: WebSocket, task_id: str):
    print("Origin:", websocket.headers.get("origin"))
    if task_id not in active_tasks:
        await websocket.close(code=1008, reason="Task not found")
        return

    await manager.connect(task_id, websocket)
    try:
        while True:
            if task_id not in active_tasks:
                break

            task = active_tasks[task_id]
            if websocket.application_state != WebSocketState.CONNECTED:
                break

            # Send progress update
            await manager.send_progress(task_id, {
                "task_id": task_id,
                "progress": task["progress"],
                "status": task["status"],
                "error": task.get("error"),
                "stage": task.get("stage"),
                "error_details": task.get("error_details")
            })
            
            # If task is completed or has error, close the connection
            if task["status"] in ["success", "error", "completed"]:

                break

            await asyncio.sleep(1)  # Keep the connection open
    except WebSocketDisconnect:
        pass
    except Exception as e:
        print(f"Unexpected WebSocket error for task {task_id}: {str(e)}")
    finally:
        if websocket.application_state == WebSocketState.CONNECTED:
            try:
                await websocket.close(code=1000, reason="Task completed")
            except Exception:
                pass
        manager.disconnect(task_id)

# Add shutdown event handler
@app.on_event("shutdown")
async def shutdown_event():
    await manager.shutdown()

# Helper function to create a new task
def create_task() -> str:
    task_id = str(uuid.uuid4())
    active_tasks[task_id] = {
        "id": task_id,
        "progress": 0,
        "status": "created",
        "result_file": None,
        "error": None,
        "stage": "created",
        "error_details": None
    }
    return task_id

# Cleanup task after completion
async def cleanup_task(task_id: str, delay: int = 600):  # 10 minutes default
    await asyncio.sleep(delay)
    if task_id in active_tasks:
        result_file = active_tasks[task_id].get("result_file")
        if result_file and os.path.exists(result_file):
            try:
                os.unlink(result_file)
            except:
                pass
        del active_tasks[task_id]

# 1. Compress PDF
@app.post("/compress_pdf/", description="Compress a PDF file")
async def compress_pdf(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    quality: int = Form(80, description="Compression quality (0-100)")
):
    task_id = create_task()
    file_bytes = await pdf_file.read()
    background_tasks.add_task(process_compress_pdf, task_id, file_bytes, quality)
    return {"task_id": task_id}

def find_ghostscript():
    candidates = ["gswin64c", "gswin32c", "gs"]

    for exe in candidates:
        path = shutil.which(exe)
        if path:
            return path

    common_paths = [
        r"C:\Program Files\gs",
        r"C:\Program Files (x86)\gs"
    ]

    for base in common_paths:
        if not os.path.exists(base):
            continue

        for version in sorted(os.listdir(base), reverse=True):
            path = os.path.join(
                base,
                version,
                "bin",
                "gswin64c.exe"
            )

            if os.path.exists(path):
                return path

    return None


async def compress_with_ghostscript(
    input_path,
    output_path,
    quality="/ebook"
):
    gs = find_ghostscript()

    if not gs:
        return False

    cmd = [
        gs,
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={quality}",
        "-dNOPAUSE",
        "-dQUIET",
        "-dBATCH",
        f"-sOutputFile={output_path}",
        input_path
    ]

    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )

        _, stderr = await proc.communicate()

        if proc.returncode != 0:
            print(stderr.decode(errors="ignore"))
            return False
    except Exception as e:
        # Ghostscript is optional; any failure here should fall back to PyMuPDF pass.
        print(f"Ghostscript compression failed: {e}")
        return False

    return os.path.exists(output_path)


def _safe_save_to_buffer(doc: fitz.Document, buffer: io.BytesIO) -> None:
    """Save PDF to memory with compatibility fallbacks across PyMuPDF versions."""
    try:
        doc.save(
            buffer,
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True
        )
        return
    except TypeError:
        pass

    try:
        doc.save(
            buffer,
            garbage=4,
            deflate=True
        )
    except TypeError:
        # Older versions may support only bare save options.
        doc.save(buffer)


def _safe_save_to_file(doc: fitz.Document, output_path: str) -> None:
    """Save PDF to file with compatibility fallbacks across PyMuPDF versions."""
    try:
        doc.save(
            output_path,
            garbage=4,
            deflate=True
        )
    except TypeError:
        doc.save(output_path)


def _safe_rewrite_images(doc: fitz.Document, quality: int) -> bool:
    """Attempt lossy image rewrite with API-compatible fallbacks."""
    if not hasattr(doc, "rewrite_images"):
        return False

    safe_quality = max(30, min(95, int(quality)))

    try:
        doc.rewrite_images(
            dpi_threshold=180,
            dpi_target=150,
            quality=safe_quality,
            lossy=True,
            lossless=True
        )
        return True
    except TypeError:
        # Signature differs on some versions.
        try:
            doc.rewrite_images(quality=safe_quality)
            return True
        except Exception:
            return False
    except Exception:
        return False

async def process_compress_pdf(task_id: str, file_bytes: bytes, quality: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("compressed", ".pdf")
    stage = "initializing"

    try:
        stage = "saving-input"
        with open(input_path, "wb") as f:
            f.write(file_bytes)

        input_size = len(file_bytes)

        await update_progress(task_id, 10, "processing", stage=stage)

        # --------------------------------------------------
        # PASS 1 - LOSSLESS OPTIMIZATION
        # --------------------------------------------------
        stage = "pass-1-lossless"
        doc = fitz.open(input_path)

        buffer = io.BytesIO()

        _safe_save_to_buffer(doc, buffer)

        optimized_size = buffer.getbuffer().nbytes

        doc.close()

        await update_progress(task_id, 25, stage=stage)

        # If already compressed enough, use it
        if optimized_size < input_size * 0.90:

            with open(output_path, "wb") as f:
                f.write(buffer.getvalue())

            active_tasks[task_id]["result_file"] = output_path

            await update_progress(
                task_id,
                100,
                "completed",
                stage="done"
            )

            return

        # --------------------------------------------------
        # PASS 2 - GHOSTSCRIPT
        # --------------------------------------------------
        stage = "pass-2-ghostscript"
        await update_progress(task_id, 40, stage=stage)

        gs_success = await compress_with_ghostscript(
            input_path,
            output_path,
            "/ebook"
        )

        await update_progress(task_id, 80, stage=stage)

        if gs_success:

            compressed_size = os.path.getsize(output_path)

            if compressed_size < input_size:

                active_tasks[task_id]["result_file"] = output_path

                await update_progress(
                    task_id,
                    100,
                    "completed",
                    stage="done"
                )

                return

        # --------------------------------------------------
        # PASS 3 - PYMUPDF FALLBACK
        # --------------------------------------------------
        stage = "pass-3-pymupdf-fallback"
        doc = fitz.open(input_path)

        if _safe_rewrite_images(doc, quality):
            try:
                _safe_save_to_file(doc, output_path)
            except Exception:
                with open(output_path, "wb") as f:
                    f.write(buffer.getvalue())

        else:

            with open(output_path, "wb") as f:
                f.write(buffer.getvalue())

        doc.close()

        # --------------------------------------------------
        # FINAL SIZE CHECK
        # --------------------------------------------------
        if os.path.exists(output_path):

            compressed_size = os.path.getsize(output_path)

            if compressed_size >= input_size:

                if optimized_size < input_size:

                    with open(output_path, "wb") as f:
                        f.write(buffer.getvalue())

                else:

                    with open(output_path, "wb") as f:
                        f.write(file_bytes)

        active_tasks[task_id]["result_file"] = output_path

        await update_progress(
            task_id,
            100,
            "completed",
            stage="done"
        )

    except Exception as e:
        error_details = traceback.format_exc()
        print(f"Compression failed at stage '{stage}': {e}")
        print(error_details)

        await update_progress(
            task_id,
            0,
            "error",
            f"{stage}: {e}",
            stage=stage,
            error_details=error_details
        )

        if os.path.exists(input_path):
            os.unlink(input_path)

        if os.path.exists(output_path):
            os.unlink(output_path)

# 2. Convert PDF to Images
@app.post("/pdf-to-images/", description="Convert PDF pages to images")
async def pdf_to_images(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    dpi: int = Form(200, description="Image DPI (150-600)"),
    format: str = Form("png", description="Image format (png, jpg, webp)"),
    page_range: Optional[str] = Form(None, description="Page range, e.g. '1-3,5,7-9'")
):
    task_id = create_task()
    background_tasks.add_task(process_pdf_to_images, task_id, pdf_file, dpi, format, page_range)
    return {"task_id": task_id}

def _parse_page_range(page_range: Optional[str], total_pages: int) -> List[int]:
    if not page_range:
        return list(range(total_pages))

    selected_pages: List[int] = []
    tokens = [token.strip() for token in page_range.split(",") if token.strip()]

    if not tokens:
        raise Exception("Invalid page range")

    for token in tokens:
        if "-" in token:
            parts = token.split("-", 1)
            if len(parts) != 2 or not parts[0].strip().isdigit() or not parts[1].strip().isdigit():
                raise Exception(f"Invalid page range segment: {token}")

            start = int(parts[0].strip())
            end = int(parts[1].strip())

            if start > end:
                raise Exception(f"Invalid page range segment: {token}")

            for page_num in range(start, end + 1):
                if page_num < 1 or page_num > total_pages:
                    raise Exception(f"Page {page_num} is out of range. Valid pages: 1-{total_pages}")
                selected_pages.append(page_num - 1)
        else:
            if not token.isdigit():
                raise Exception(f"Invalid page number: {token}")
            page_num = int(token)
            if page_num < 1 or page_num > total_pages:
                raise Exception(f"Page {page_num} is out of range. Valid pages: 1-{total_pages}")
            selected_pages.append(page_num - 1)

    # Keep order while removing duplicates.
    deduped_pages = list(dict.fromkeys(selected_pages))
    if not deduped_pages:
        raise Exception("No valid pages selected")

    return deduped_pages


async def process_pdf_to_images(task_id: str, pdf_file: UploadFile, dpi: int, format: str, page_range: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("images", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("images", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        pages_to_convert = _parse_page_range(page_range, total_pages)
        
        # Map file extension to Pillow save format names.
        format_map = {
            "png": "PNG",
            "jpg": "JPEG",
            "jpeg": "JPEG",
            "tiff": "TIFF",
            "bmp": "BMP",
            "webp": "WEBP",
        }
        format_key = (format or "png").strip().lower()
        pil_format = format_map.get(format_key)
        if not pil_format:
            raise Exception(f"Unsupported image format: {format}")

        # Convert each page to image
        image_files = []
        total_selected_pages = len(pages_to_convert)
        for i, page_idx in enumerate(pages_to_convert):
            page = pdf_document[page_idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
            
            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Save as requested format
            img_path = os.path.join(output_dir, f"page_{page_idx + 1}.{format_key}")
            img.save(img_path, format=pil_format)
            image_files.append(img_path)
            
            # Update progress
            progress = 10 + int(80 * (i + 1) / total_selected_pages)
            await update_progress(task_id, progress)
        
        # Create ZIP archive with all images
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in image_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual image files
        for file in image_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        print(output_zip)
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 3. Convert Images to PDF
@app.post("/images-to-pdf/", description="Convert multiple images to a single PDF")
async def images_to_pdf(
    background_tasks: BackgroundTasks,
    images: List[UploadFile] = File(...),
    page_size: PageSize = Form(PageSize.A4),
    margin: int = Form(50, description="Margin in pixels")
):
    task_id = create_task()
    background_tasks.add_task(process_images_to_pdf, task_id, images, page_size, margin)
    return {"task_id": task_id}

async def process_images_to_pdf(task_id: str, images: List[UploadFile], page_size: PageSize, margin: int):
    output_path = get_temp_file("converted", ".pdf")
    temp_images = []
    
    try:
        # Page size dimensions (width, height) in points
        page_sizes = {
            PageSize.A4: (595, 842),
            PageSize.LETTER: (612, 792),
            PageSize.LEGAL: (612, 1008),
            PageSize.TABLOID: (792, 1224)
        }
        page_width, page_height = page_sizes[page_size]
        
        # Create new PDF
        pdf = fitz.open()
        
        total_images = len(images)
        for i, img_file in enumerate(images):
            # Save image to temp file
            img_path = get_temp_file(f"img_{i}", f".{img_file.filename.split('.')[-1]}")
            with open(img_path, "wb") as f:
                content = await img_file.read()
                f.write(content)
            temp_images.append(img_path)
            
            # Open image
            img = Image.open(img_path)
            
            # Calculate scaling to fit within page with margins
            img_width, img_height = img.size
            max_width = page_width - 2 * margin
            max_height = page_height - 2 * margin
            
            scale = min(max_width / img_width, max_height / img_height)
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            
            # Create new page
            page = pdf.new_page(width=page_width, height=page_height)
            
            # Calculate position to center image
            x = (page_width - new_width) / 2
            y = (page_height - new_height) / 2
            
            # Insert image
            page.insert_image(fitz.Rect(x, y, x + new_width, y + new_height), filename=img_path)
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_images)
            await update_progress(task_id, progress)
        
        # Save PDF
        pdf.save(output_path)
        pdf.close()
        
        # Cleanup temp images
        for img_path in temp_images:
            os.unlink(img_path)
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        for img_path in temp_images:
            if os.path.exists(img_path):
                os.unlink(img_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 4. Extract Text from PDF
@app.post("/extract-text/", description="Extract text from a PDF file")
async def extract_text(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_range: Optional[str] = Form(None, description="Page range (e.g., 1-5)")
):
    task_id = create_task()
    background_tasks.add_task(process_extract_text, task_id, pdf_file, page_range)
    return {"task_id": task_id}

async def process_extract_text(task_id: str, pdf_file: UploadFile, page_range: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("extracted_text", ".txt")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Parse page range
        if page_range:
            parts = page_range.split('-')
            start_page = max(0, int(parts[0]) - 1)  # Convert from 1-based to 0-based
            end_page = min(total_pages, int(parts[1]) if len(parts) > 1 else start_page + 1)
        else:
            start_page = 0
            end_page = total_pages
        
        # Extract text from pages
        with open(output_path, "w", encoding="utf-8") as f:
            for i in range(start_page, end_page):
                page = pdf_document[i]
                text = page.get_text()
                f.write(f"--- Page {i+1} ---\n\n")
                f.write(text)
                f.write("\n\n")
                
                # Update progress
                progress = 10 + int(85 * (i - start_page + 1) / (end_page - start_page))
                await update_progress(task_id, progress)
        
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 5. Extract Images from PDF
@app.post("/extract-images/", description="Extract images from a PDF file")
async def extract_images(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    min_size: int = Form(100, description="Minimum image size in pixels")
):
    task_id = create_task()
    background_tasks.add_task(process_extract_images, task_id, pdf_file, min_size)
    return {"task_id": task_id}

async def process_extract_images(task_id: str, pdf_file: UploadFile, min_size: int):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("images", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("extracted_images", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Extract images from each page
        image_files = []
        image_count = 0
        
        for i, page in enumerate(pdf_document):
            # Get image list
            image_list = page.get_images(full=True)
            
            # Process each image
            for j, img in enumerate(image_list):
                xref = img[0]  # image reference
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Load as PIL Image to check size
                try:
                    img = Image.open(io.BytesIO(image_bytes))
                    if min(img.size) < min_size:
                        continue  # Skip small images
                    
                    # Save image
                    img_path = os.path.join(output_dir, f"image_{image_count}.{base_image['ext']}")
                    with open(img_path, "wb") as f:
                        f.write(image_bytes)
                    image_files.append(img_path)
                    image_count += 1
                except Exception as e:
                    # Skip problematic images
                    continue
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Create ZIP archive with all images
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in image_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual image files
        for file in image_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 6. Merge PDFs
@app.post("/merge-pdfs/", description="Merge multiple PDF files into one")
async def merge_pdfs(
    background_tasks: BackgroundTasks,
    pdf_files: List[UploadFile] = File(...)
):
    task_id = create_task()
    background_tasks.add_task(process_merge_pdfs, task_id, pdf_files)
    return {"task_id": task_id}

async def process_merge_pdfs(task_id: str, pdf_files: List[UploadFile]):
    temp_files = []
    output_path = get_temp_file("merged", ".pdf")
    
    try:
        await update_progress(task_id, 5, "processing")
        
        # Create merged PDF
        merged_pdf = fitz.open()
        
        total_files = len(pdf_files)
        for i, pdf_file in enumerate(pdf_files):
            # Save uploaded file
            input_path = get_temp_file(f"input_{i}", ".pdf")
            with open(input_path, "wb") as f:
                content = await pdf_file.read()
                f.write(content)
            temp_files.append(input_path)
            
            # Open PDF
            pdf_document = fitz.open(input_path)
            
            # Append to merged PDF
            merged_pdf.insert_pdf(pdf_document)
            pdf_document.close()
            
            # Update progress
            progress = 5 + int(90 * (i + 1) / total_files)
            await update_progress(task_id, progress)
        
        # Save merged PDF
        merged_pdf.save(output_path)
        merged_pdf.close()
        
        # Cleanup temp files
        for file in temp_files:
            os.unlink(file)
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        for file in temp_files:
            if os.path.exists(file):
                os.unlink(file)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 7. Split PDF
@app.post("/split-pdf/", description="Split a PDF into multiple files")
async def split_pdf(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    split_method: str = Form("pages", description="Split method: 'pages' or 'ranges'"),
    pages_per_pdf: Optional[int] = Form(None, description="Pages per output PDF"),
    page_ranges: Optional[str] = Form(None, description="Page ranges (e.g., '1-3,4-6,7-9')")
):
    task_id = create_task()
    background_tasks.add_task(process_split_pdf, task_id, pdf_file, split_method, pages_per_pdf, page_ranges)
    return {"task_id": task_id}

async def process_split_pdf(task_id: str, pdf_file: UploadFile, split_method: str, pages_per_pdf: Optional[int], page_ranges: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_dir = get_temp_file("split", "")
    os.makedirs(output_dir, exist_ok=True)
    output_zip = get_temp_file("split_pdfs", ".zip")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Split based on method
        output_files = []
        
        if split_method == "pages" and pages_per_pdf:
            # Split by number of pages per PDF
            for i in range(0, total_pages, pages_per_pdf):
                # Calculate end page
                end = min(i + pages_per_pdf, total_pages)
                
                # Create new PDF
                output_file = os.path.join(output_dir, f"split_{i+1}-{end}.pdf")
                new_pdf = fitz.open()
                
                # Add pages from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=i, to_page=end-1)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + pages_per_pdf) / total_pages)
                await update_progress(task_id, min(95, progress))
        
        elif split_method == "ranges" and page_ranges:
            # Split by specified page ranges
            ranges = page_ranges.split(',')
            
            for i, range_str in enumerate(ranges):
                parts = range_str.strip().split('-')
                
                # Parse start and end pages (convert to 0-based)
                start = max(0, int(parts[0]) - 1)
                end = min(total_pages, int(parts[1]) if len(parts) > 1 else start + 1)
                
                # Create new PDF
                output_file = os.path.join(output_dir, f"split_{start+1}-{end}.pdf")
                new_pdf = fitz.open()
                
                # Add pages from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=start, to_page=end-1)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / len(ranges))
                await update_progress(task_id, min(95, progress))
        
        else:
            # Split into individual pages
            for i in range(total_pages):
                # Create new PDF with single page
                output_file = os.path.join(output_dir, f"page_{i+1}.pdf")
                new_pdf = fitz.open()
                
                # Add page from original PDF
                new_pdf.insert_pdf(pdf_document, from_page=i, to_page=i)
                
                # Save new PDF
                new_pdf.save(output_file)
                new_pdf.close()
                output_files.append(output_file)
                
                # Update progress
                progress = 10 + int(85 * (i + 1) / total_pages)
                await update_progress(task_id, min(95, progress))
        
        pdf_document.close()
        
        # Create ZIP archive with all split PDFs
        import zipfile
        with zipfile.ZipFile(output_zip, 'w') as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))
        
        # Cleanup individual PDF files
        for file in output_files:
            os.unlink(file)
        os.rmdir(output_dir)
        
        active_tasks[task_id]["result_file"] = output_zip
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_zip):
            os.unlink(output_zip)

# 8. Add Password Protection
@app.post("/add-password/", description="Add password protection to a PDF")
async def add_password(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    owner_password: str = Form(...),
    user_password: Optional[str] = Form(None)
):
    task_id = create_task()
    password = PdfPasswordModel(
        owner_password=owner_password,
        user_password=user_password
    )
    background_tasks.add_task(process_add_password, task_id, pdf_file, password)
    return {"task_id": task_id}

async def process_add_password(task_id: str, pdf_file: UploadFile, password: PdfPasswordModel):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("protected", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 20, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        
        # Set permissions
        perm = int(
            fitz.PDF_PERM_PRINT |  # allow printing
            fitz.PDF_PERM_COPY |   # allow copying
            fitz.PDF_PERM_ANNOTATE # allow annotations
        )
        
        pdf_document.save(
            output_path,
            encryption=fitz.PDF_ENCRYPT_AES_256,  # strongest encryption
            owner_pw=password.owner_password,
            user_pw=password.user_password or "",
            permissions=perm
        )
        
        pdf_document.close()
        
        await update_progress(task_id, 90, "processing")
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 9. Remove Password Protection
@app.post("/remove-password/", description="Remove password protection from a PDF")
async def remove_password(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    password: str = Form(...)
):
    task_id = create_task()
    background_tasks.add_task(process_remove_password, task_id, pdf_file, password)
    return {"task_id": task_id}

async def process_remove_password(task_id: str, pdf_file: UploadFile, password: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("unprotected", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 20, "processing")
        
        # Open PDF with password
        pdf_document = fitz.open(input_path)
        
        # Check if PDF is encrypted
        if pdf_document.is_encrypted:
            # Try to authenticate with password
            if not pdf_document.authenticate(password):
                raise Exception("Incorrect password")
            
            # Save without encryption
            pdf_document.save(output_path)
        else:
            # No password to remove, just copy
            pdf_document.save(output_path)
        
        pdf_document.close()
        
        await update_progress(task_id, 90, "processing")
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 10. Add Watermark
@app.post("/add-watermark/", description="Add text watermark to a PDF")
async def add_watermark(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    watermark_text: str = Form(...),
    position: WatermarkPosition = Form(WatermarkPosition.CENTER),
    opacity: float = Form(0.3, description="Opacity (0.1-1.0)"),
    rotation: int = Form(45, description="Rotation angle in degrees")
):
    task_id = create_task()
    background_tasks.add_task(process_add_watermark, task_id, pdf_file, watermark_text, position, opacity, rotation)
    return {"task_id": task_id}

async def process_add_watermark(task_id: str, pdf_file: UploadFile, watermark_text: str, position: WatermarkPosition, opacity: float, rotation: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("watermarked", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Process each page
        for i, page in enumerate(pdf_document):
            # Create watermark with proper rotation and positioning
            font_size = min(page.rect.width, page.rect.height) / 10
            
            # Determine position for watermark
            if position == WatermarkPosition.CENTER:
                x = page.rect.width / 2
                y = page.rect.height / 2
            elif position == WatermarkPosition.TOP_LEFT:
                x = page.rect.width * 0.2
                y = page.rect.height * 0.2
            elif position == WatermarkPosition.TOP_RIGHT:
                x = page.rect.width * 0.8
                y = page.rect.height * 0.2
            elif position == WatermarkPosition.BOTTOM_LEFT:
                x = page.rect.width * 0.2
                y = page.rect.height * 0.8
            elif position == WatermarkPosition.BOTTOM_RIGHT:
                x = page.rect.width * 0.8
                y = page.rect.height * 0.8
            
            # Insert watermark text
            page.insert_text(
                (x, y),
                watermark_text,
                fontsize=font_size,
                rotate=rotation,
                color=(0, 0, 0),
                fill_opacity=max(0.0, min(1.0, opacity)),
                render_mode=2  # Fill with transparency
            )
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 11. Edit Metadata
@app.post("/edit-metadata/", description="Edit PDF metadata")
async def edit_metadata(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    author: Optional[str] = Form(None),
    subject: Optional[str] = Form(None),
    keywords: Optional[str] = Form(None),
    creator: Optional[str] = Form(None),
    producer: Optional[str] = Form(None),
):
    task_id = create_task()
    background_tasks.add_task(
        process_edit_metadata,
        task_id,
        pdf_file,
        title,
        author,
        subject,
        keywords,
        creator,
        producer,
    )
    return {"task_id": task_id}


async def process_edit_metadata(
    task_id: str,
    pdf_file: UploadFile,
    title: Optional[str],
    author: Optional[str],
    subject: Optional[str],
    keywords: Optional[str],
    creator: Optional[str],
    producer: Optional[str],
):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("metadata-edited", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(await pdf_file.read())

        await update_progress(task_id, 20, "processing")

        pdf_document = fitz.open(input_path)
        existing_metadata = pdf_document.metadata or {}

        editable_fields = (
            "title",
            "author",
            "subject",
            "keywords",
            "creator",
            "producer",
        )
        metadata = {key: str(existing_metadata.get(key, "") or "") for key in editable_fields}
        metadata.update({
            "title": title if title is not None else metadata.get("title", ""),
            "author": author if author is not None else metadata.get("author", ""),
            "subject": subject if subject is not None else metadata.get("subject", ""),
            "keywords": keywords if keywords is not None else metadata.get("keywords", ""),
            "creator": creator if creator is not None else metadata.get("creator", ""),
            "producer": producer if producer is not None else metadata.get("producer", ""),
        })

        pdf_document.set_metadata(metadata)

        try:
            if hasattr(pdf_document, "get_xml_metadata") and pdf_document.get_xml_metadata():
                pdf_document.del_xml_metadata()
        except Exception:
            pass

        pdf_document.save(output_path, garbage=4, deflate=True)
        pdf_document.close()

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)


def _parse_remove_page_ranges(page_ranges: str, total_pages: int) -> List[int]:
    selected_pages: List[int] = []

    for token in [part.strip() for part in (page_ranges or "").split(",") if part.strip()]:
        if "-" in token:
            start_text, end_text = [part.strip() for part in token.split("-", 1)]
            if not start_text.isdigit() or not end_text.isdigit():
                raise ValueError(f"Invalid page range: {token}")

            start_page = int(start_text)
            end_page = int(end_text)
            if start_page < 1 or end_page < start_page:
                raise ValueError(f"Invalid page range: {token}")

            for page_number in range(start_page, end_page + 1):
                if page_number > total_pages:
                    raise ValueError(f"Page {page_number} is out of range")
                selected_pages.append(page_number - 1)
        else:
            if not token.isdigit():
                raise ValueError(f"Invalid page number: {token}")

            page_number = int(token)
            if page_number < 1 or page_number > total_pages:
                raise ValueError(f"Page {page_number} is out of range")
            selected_pages.append(page_number - 1)

    unique_pages = list(dict.fromkeys(selected_pages))
    if not unique_pages:
        raise ValueError("No pages were selected")

    return unique_pages


# 12. Remove Pages
@app.post("/remove-pages/", description="Remove selected pages from a PDF")
async def remove_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_ranges: str = Form(..., description="Comma-separated pages/ranges to remove, e.g. '1,3,5-7'"),
):
    task_id = create_task()
    background_tasks.add_task(process_remove_pages, task_id, pdf_file, page_ranges)
    return {"task_id": task_id}


async def process_remove_pages(task_id: str, pdf_file: UploadFile, page_ranges: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("pages-removed", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(await pdf_file.read())

        await update_progress(task_id, 20, "processing")

        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        pages_to_remove = _parse_remove_page_ranges(page_ranges, total_pages)
        pages_to_keep = [page_index for page_index in range(total_pages) if page_index not in pages_to_remove]

        if not pages_to_keep:
            raise ValueError("At least one page must remain in the PDF")

        pdf_document.select(pages_to_keep)
        await update_progress(task_id, 75, "processing")

        pdf_document.save(output_path, garbage=4, deflate=True)
        pdf_document.close()

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 13. Rotate Pages
@app.post("/rotate-pages/", description="Rotate pages in a PDF")
async def rotate_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    angle: Rotation = Form(Rotation.DEGREES_90),
    page_range: Optional[str] = Form(None, description="Page range (e.g., 1-5)")
):
    task_id = create_task()
    background_tasks.add_task(process_rotate_pages, task_id, pdf_file, angle, page_range)
    return {"task_id": task_id}

async def process_rotate_pages(task_id: str, pdf_file: UploadFile, angle: Rotation, page_range: Optional[str]):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("rotated", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Parse page range
        if page_range:
            parts = page_range.split('-')
            start_page = max(0, int(parts[0]) - 1)  # Convert from 1-based to 0-based
            end_page = min(total_pages, int(parts[1]) if len(parts) > 1 else start_page + 1)
        else:
            start_page = 0
            end_page = total_pages
        
        # Rotate pages
        for i in range(start_page, end_page):
            pdf_document[i].set_rotation(angle)
            
            # Update progress
            progress = 10 + int(85 * (i - start_page + 1) / (end_page - start_page))
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 12. Reorder Pages
@app.post("/reorder-pages/", description="Reorder pages in a PDF")
async def reorder_pages(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    page_order: str = Form(..., description="Comma-separated page numbers (e.g., '3,1,2,4')")
):
    task_id = create_task()
    background_tasks.add_task(process_reorder_pages, task_id, pdf_file, page_order)
    return {"task_id": task_id}

async def process_reorder_pages(task_id: str, pdf_file: UploadFile, page_order: str):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("reordered", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Parse page order (convert from 1-based to 0-based)
        try:
            page_numbers = [int(p.strip()) - 1 for p in page_order.split(',')]
            
            # Validate page numbers
            if any(p < 0 or p >= total_pages for p in page_numbers):
                raise Exception(f"Invalid page numbers. Must be between 1 and {total_pages}.")
        except ValueError:
            raise Exception("Invalid page order format. Use comma-separated page numbers.")
        
        # Create new PDF with reordered pages
        new_pdf = fitz.open()
        
        for i, page_num in enumerate(page_numbers):
            new_pdf.insert_pdf(pdf_document, from_page=page_num, to_page=page_num)
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / len(page_numbers))
            await update_progress(task_id, progress)
        
        # Save result
        new_pdf.save(output_path)
        new_pdf.close()
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 13. Add Page Numbers
@app.post("/add-page-numbers/", description="Add page numbers to a PDF")
async def add_page_numbers(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    position: str = Form("bottom-right", description="Position: 'bottom-right', 'bottom-center', 'bottom-left', 'top-right', 'top-center', 'top-left'"),
    start_number: int = Form(1, description="Starting page number"),
    font_size: int = Form(12, description="Font size")
):
    task_id = create_task()
    background_tasks.add_task(process_add_page_numbers, task_id, pdf_file, position, start_number, font_size)
    return {"task_id": task_id}

async def process_add_page_numbers(task_id: str, pdf_file: UploadFile, position: str, start_number: int, font_size: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("page_numbers", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Add page numbers to each page
        for i in range(total_pages):
            page = pdf_document[i]
            page_number = start_number + i
            
            # Determine position
            margin = 30  # margin in points
            
            if position == "bottom-right":
                x = page.rect.width - margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_RIGHT
            elif position == "bottom-center":
                x = page.rect.width / 2
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_CENTER
            elif position == "bottom-left":
                x = margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_LEFT
            elif position == "top-right":
                x = page.rect.width - margin
                y = margin
                align = fitz.TEXT_ALIGN_RIGHT
            elif position == "top-center":
                x = page.rect.width / 2
                y = margin
                align = fitz.TEXT_ALIGN_CENTER
            elif position == "top-left":
                x = margin
                y = margin
                align = fitz.TEXT_ALIGN_LEFT
            else:
                # Default to bottom right
                x = page.rect.width - margin
                y = page.rect.height - margin
                align = fitz.TEXT_ALIGN_RIGHT
            
            # Insert page number
            page.insert_text(
                (x, y),
                str(page_number),
                fontsize=font_size,
                align=align
            )
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 14. Add Headers/Footers
@app.post("/add-header-footer/", description="Add headers and/or footers to a PDF")
async def add_header_footer(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    header_text: Optional[str] = Form(None, description="Header text"),
    footer_text: Optional[str] = Form(None, description="Footer text"),
    font_size: int = Form(10, description="Font size")
):
    task_id = create_task()
    background_tasks.add_task(process_add_header_footer, task_id, pdf_file, header_text, footer_text, font_size)
    return {"task_id": task_id}

async def process_add_header_footer(task_id: str, pdf_file: UploadFile, header_text: Optional[str], footer_text: Optional[str], font_size: int):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("header_footer", ".pdf")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            content = await pdf_file.read()
            f.write(content)
        
        await update_progress(task_id, 10, "processing")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        
        # Add header/footer to each page
        for i in range(total_pages):
            page = pdf_document[i]
            
            # Add header if specified
            if header_text:
                page.insert_text(
                    (page.rect.width / 2, 20),  # Position at top center
                    header_text,
                    fontsize=font_size,
                    align=fitz.TEXT_ALIGN_CENTER
                )
            
            # Add footer if specified
            if footer_text:
                page.insert_text(
                    (page.rect.width / 2, page.rect.height - 20),  # Position at bottom center
                    footer_text,
                    fontsize=font_size,
                    align=fitz.TEXT_ALIGN_CENTER
                )
            
            # Update progress
            progress = 10 + int(85 * (i + 1) / total_pages)
            await update_progress(task_id, progress)
        
        # Save result
        pdf_document.save(output_path)
        pdf_document.close()
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        await update_progress(task_id, 0, "error", str(e))
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# 15. PDF to Word conversion
@app.post("/pdf_to_word/", description="Convert PDF to Word format (Note: Basic conversion only)")
async def pdf_to_word(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...)
):
    task_id = create_task()
    file_bytes = await pdf_file.read()
    background_tasks.add_task(process_pdf_to_word, task_id, file_bytes)
    return {"task_id": task_id}


def _convert_pdf_to_docx_with_aspose(input_path: str, output_path: str) -> None:
    """Convert PDF to DOCX using Aspose.PDF."""
    pdf_document = aspose_pdf.Document(input_path)
    pdf_document.save(output_path, aspose_pdf.SaveFormat.DOC_X)


def _extract_pdf_to_docx_with_alignment(input_path: str, output_path: str) -> int:
    """Fallback converter preserving basic paragraph alignment from PDF blocks."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pdf_document = fitz.open(input_path)
    total_pages = len(pdf_document)
    doc = Document()

    for page_index, page in enumerate(pdf_document):
        page_width = page.rect.width
        blocks = page.get_text("blocks", sort=True)

        for block in blocks:
            # block format: (x0, y0, x1, y1, text, block_no, block_type)
            if len(block) < 5:
                continue

            x0, _, x1, _, text = block[:5]
            if not isinstance(text, str):
                continue

            text = text.strip()
            if not text:
                continue

            paragraph = doc.add_paragraph(text)
            center_x = (x0 + x1) / 2.0
            block_width = x1 - x0

            # Heuristics for common left/center/right aligned blocks.
            if abs(center_x - page_width / 2.0) <= page_width * 0.10 and block_width <= page_width * 0.70:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif x0 >= page_width * 0.55:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if page_index < total_pages - 1:
            doc.add_page_break()

    doc.save(output_path)
    pdf_document.close()
    return total_pages

async def process_pdf_to_word(task_id: str, pdf_file: bytes):
    input_path = get_temp_file("input", ".pdf")
    output_path = get_temp_file("output", ".docx")
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as f:
            f.write(pdf_file)
        
        await update_progress(task_id, 5, stage="pdf-to-word-init")

        pdf_document = fitz.open(input_path)
        total_pages = len(pdf_document)
        pdf_document.close()

        await update_progress(task_id, 10, stage="pdf-to-word-prepare")

        try:
            await update_progress(task_id, 15, stage="aspose-converting")

            loop = asyncio.get_event_loop()
            conversion_future = loop.run_in_executor(
                None,
                lambda: _convert_pdf_to_docx_with_aspose(input_path, output_path)
            )

            heartbeat_progress = 15
            while not conversion_future.done():
                heartbeat_progress = min(90, heartbeat_progress + 2)
                await update_progress(task_id, heartbeat_progress, stage="aspose-converting")
                await asyncio.sleep(1)

            await conversion_future

            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise Exception("Aspose conversion completed but no output DOCX was generated")

            await update_progress(task_id, 95, stage="aspose-converting")

        except Exception as convert_error:
            print(f"Aspose conversion failed, using alignment fallback: {convert_error}")
            await update_progress(task_id, 25, stage="fallback-aligned-extraction")

            fallback_pages = _extract_pdf_to_docx_with_alignment(input_path, output_path)
            for i in range(max(1, fallback_pages)):
                progress = 25 + int(70 * (i + 1) / max(1, fallback_pages))
                await update_progress(task_id, progress, stage="fallback-aligned-extraction")
        
        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed", stage="done")
        
        # Schedule cleanup
        asyncio.create_task(cleanup_task(task_id))
    
    except Exception as e:
        error_details = traceback.format_exc()
        await update_progress(
            task_id,
            0,
            "error",
            str(e),
            stage="pdf-to-word-error",
            error_details=error_details
        )
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)


# 16. Word to PDF conversion
@app.post("/word_to_pdf/", description="Convert Word documents to PDF format")
async def word_to_pdf(
    background_tasks: BackgroundTasks,
    word_file: UploadFile = File(...)
):
    task_id = create_task()
    file_bytes = await word_file.read()
    background_tasks.add_task(process_word_to_pdf, task_id, file_bytes, word_file.filename or "document")
    return {"task_id": task_id}


async def process_word_to_pdf(task_id: str, file_bytes: bytes, original_filename: str):
    input_path = get_temp_file("input", ".docx")
    output_path = get_temp_file("output", ".pdf")

    try:
        with open(input_path, "wb") as f:
            f.write(file_bytes)

        await update_progress(task_id, 5, stage="word-to-pdf-init")

        try:
            from docx import Document
        except Exception as import_error:
            raise Exception(f"python-docx is required for Word to PDF conversion: {import_error}")

        document = Document(input_path)
        paragraph_count = max(1, len(document.paragraphs))

        _convert_docx_to_pdf(input_path, output_path)

        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise Exception("Word to PDF conversion did not generate a PDF")

        await update_progress(task_id, 95, stage="word-to-pdf-complete")

        active_tasks[task_id]["result_file"] = output_path
        await update_progress(task_id, 100, "completed", stage="done")

        asyncio.create_task(cleanup_task(task_id))

    except Exception as e:
        error_details = traceback.format_exc()
        await update_progress(
            task_id,
            0,
            "error",
            str(e),
            stage="word-to-pdf-error",
            error_details=error_details,
        )
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

# Endpoint to get task status
@app.get("/task/{task_id}", description="Get the status of a task")
async def get_task_status(task_id: str):
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    return active_tasks[task_id]

# Endpoint to download the result file
@app.get("/download/{task_id}", description="Download the result file")
async def download_result(task_id: str):

    #filepath = build_temp_file(task_id)
    #if not os.path.exists(filepath):
    #    raise HTTPException(status_code=404, detail="Task not found")

    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Task not completed")
    
    if not task["result_file"] or not os.path.exists(task["result_file"]):
        raise HTTPException(status_code=404, detail="Result file not found")
    
    # Get original filename from the path
    filename = os.path.basename(task["result_file"])

    return FileResponse(
        path=task["result_file"], 
        filename=filename,
        media_type="application/octet-stream"  # Force download
    )

@app.post("/api/execute", response_model=ApiResponse)
async def execute_request(request: ApiRequest):
    start_time = time.time()
    request_id = str(uuid.uuid4())
    
    # Replace environment variables if they exist
    # This would be implemented to handle {{variable}} format
    
    # Build URL with query parameters
    url = request.url
    if request.query_params:
        # Append query params to URL
        if "?" not in url:
            url += "?"
        else:
            if not url.endswith("&") and not url.endswith("?"):
                url += "&"
        
        query_parts = []
        for key, value in request.query_params.items():
            if value:  # Only add if value is not empty
                query_parts.append(f"{key}={value}")
        
        url += "&".join(query_parts)
    
    # Prepare request
    headers = request.headers or {}
    
    # Handle authentication
    if request.auth:
        auth_type = request.auth.get("type")
        if auth_type == "basic":
            # Implement basic auth
            pass
        elif auth_type == "bearer":
            headers["Authorization"] = f"Bearer {request.auth.get('token', '')}"
        elif auth_type == "api_key":
            key_name = request.auth.get("key_name", "")
            key_value = request.auth.get("key_value", "")
            if request.auth.get("in") == "header":
                headers[key_name] = key_value
            else:  # in query
                if "?" not in url:
                    url += f"?{key_name}={key_value}"
                else:
                    url += f"&{key_name}={key_value}"
    
    # Execute the request
    async with httpx.AsyncClient(follow_redirects=True) as client:
        try:
            # Handle different body types
            data = None
            files = None
            json_data = None
            
            if request.body_type == BodyType.RAW and request.body:
                if request.raw_body_format == RawBodyFormat.JSON:
                    try:
                        json_data = request.body
                    except:
                        json_data = request.body  # Already in correct format
                else:
                    # Handle other raw formats (TEXT, XML, HTML)
                    data = request.body
            
            elif request.body_type == BodyType.FORM_DATA and request.form_data:
                # In a real implementation, this would handle file uploads too
                data = request.form_data
            
            elif request.body_type == BodyType.URLENCODED and request.form_data:
                data = request.form_data
            
            # Make the request
            response = await client.request(
                method=request.method,
                url=url,
                headers=headers,
                data=data,
                files=files,
                json=json_data,
                timeout=30.0  # 30 second timeout
            )
            
            # Calculate time taken
            time_taken = (time.time() - start_time) * 1000  # convert to ms
            
            # Convert headers to dict (httpx returns a Headers object)
            response_headers = dict(response.headers.items())
            
            # Get response size
            response_size = len(response.content)
            
            # Try to parse response body based on content type
            try:
                if "application/json" in response.headers.get("content-type", "").lower():
                    response_body = response.json()
                else:
                    response_body = response.text
            except:
                response_body = response.text
            
            # Get cookies
            cookies = dict(response.cookies)
            
            # Store in history
            api_response = ApiResponse(
                request_id=request_id,
                request=request,
                status_code=response.status_code,
                headers=response_headers,
                body=response_body,
                cookies=cookies,
                time_taken=time_taken,
                size=response_size,
                timestamp=datetime.now()
            )
            
            db.requests_history.append(api_response)
            
            return api_response
            
        except httpx.RequestError as e:
            raise HTTPException(status_code=500, detail=f"Request error: {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.get("/api/history", response_model=List[ApiResponse])
async def get_history():
    return db.requests_history

@app.post("/api/collections", response_model=Collection)
async def create_collection(collection: Collection):
    if not collection.id:
        collection.id = str(uuid.uuid4())
    db.collections[collection.id] = collection
    return collection

@app.get("/api/collections", response_model=List[Collection])
async def get_collections():
    return list(db.collections.values())

@app.get("/api/collections/{collection_id}", response_model=Collection)
async def get_collection(collection_id: str):
    if collection_id not in db.collections:
        raise HTTPException(status_code=404, detail="Collection not found")
    return db.collections[collection_id]

@app.post("/api/collections/{collection_id}/requests", response_model=SavedRequest)
async def add_request_to_collection(collection_id: str, request: SavedRequest):
    if collection_id not in db.collections:
        raise HTTPException(status_code=404, detail="Collection not found")
    
    if not request.id:
        request.id = str(uuid.uuid4())
    
    db.collections[collection_id].requests.append(request)
    return request

@app.post("/api/environments", response_model=Environment)
async def create_environment(environment: Environment):
    if not environment.id:
        environment.id = str(uuid.uuid4())
    db.environments[environment.id] = environment
    return environment

@app.get("/api/environments", response_model=List[Environment])
async def get_environments():
    return list(db.environments.values())

@app.get("/api/environments/{environment_id}", response_model=Environment)
async def get_environment(environment_id: str):
    if environment_id not in db.environments:
        raise HTTPException(status_code=404, detail="Environment not found")
    return db.environments[environment_id]

@app.put("/api/environments/{environment_id}", response_model=Environment)
async def update_environment(environment_id: str, environment: Environment):
    if environment_id not in db.environments:
        raise HTTPException(status_code=404, detail="Environment not found")
    
    environment.id = environment_id
    db.environments[environment_id] = environment
    return environment

@app.post("/youtube/info")
async def showThumbnail(request: Request):
    try:
        body = await request.json()

        youtube_url = body.get("url")
        if not youtube_url:
            raise HTTPException(status_code=400, detail="Missing 'url' in request body")

        ydl_opts = {
            'format': 'best',  # Download the best quality available
            'outtmpl': 'test.mp4',  # Save file with the video title as the filename
        }
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=False)
            #print(info.get('formats', []))
            #print(info.get('formats', []))

            def safe_format(format_entry):
                print(format_entry.get('acodec'))
                return {
                    key: value for key, value in {
                        'format_id': format_entry.get('format_id'),
                        'ext': format_entry.get('ext'),
                        'resolution': format_entry.get('resolution'),
                        'fps': format_entry.get('fps'),
                        'filesize': format_entry.get('filesize'),
                        'format_note': format_entry.get('format_note'),
                        'height': format_entry.get('height'),
                        'width': format_entry.get('width'),
                        'url': format_entry.get('url'),
                        'acodec': format_entry.get('acodec'),
                        'vcodec': format_entry.get('vcodec'),
                    }.items() if value is not None
                }
            video_info = {
                'id': info.get('id'),
                'title': info.get('title'),
                'thumbnail': info.get('thumbnail'),
                'duration': info.get('duration'),
                'formats': [safe_format(f) for f in info.get('formats', []) if f.get('acodec') not in ['none', None]],
                'upload_date': info.get('upload_date'),
                'uploader': info.get('uploader'),
                'view_count': info.get('view_count'),
            }

    except HTTPException:
        raise
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=f"Failed to fetch YouTube info: {e}")
    return video_info

@app.post("/api/sitemap/analyze")
async def analyze_sitemap_main(url: str, max_depth: Optional[int] = 3):
    return await analyze_sitemap(url, max_depth)

@app.post("/svg-convert")
async def convert_image_to_svg_main(
    image_file: UploadFile = File(...),
    threshold: int = Form(128, description="Threshold (0-255)"),
    simplify: float = Form(2.0, description="Simplification level (0-10)"),
    smoothing: int = Form(5, description="Smoothing level (0-10)"),
    edge_detection: bool = Form(False, description="Apply edge detection"),
):
    return await convert_image_to_svg(image_file=image_file, threshold=threshold, simplify=simplify,
                                      smoothing=smoothing, edge_detection=edge_detection)

@app.get("/download/{filename}")
async def download_svg(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(file_path, headers={
        "Content-Disposition": f"attachment; filename={filename}"
    })

@app.get("/qr-styles")
async def get_qrstyles():
   return await get_styles()

@app.post("/generate-qr")
async def generate_qr(
    data: str = Form(...),
    size: int = Form(10),
    border: int = Form(4),
    error_correction: str = Form("M"),
    fill_color: str = Form("#000000"),
    back_color: str = Form("#FFFFFF"),
    module_drawer: str = Form("square"),
    color_mask: str = Form("solid"),
    gradient_start: Optional[str] = Form(None),
    gradient_end: Optional[str] = Form(None),
    logo: Optional[UploadFile] = File(None),
):
    return await generate_qr_code(data=data, size=size, border=border, error_correction=error_correction,
                     fill_color=fill_color, back_color=back_color, module_drawer=module_drawer,
                     color_mask=color_mask, gradient_start=gradient_start, gradient_end=gradient_end,
                     logo=logo)

# Health check endpoint
@app.get("/health", description="Health check endpoint")
async def health_check():
    return {"status": "healthy"}

# Run the server when executed directly
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8001, reload=True)